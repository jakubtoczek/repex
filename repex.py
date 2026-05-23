#!/usr/bin/env python3
r"""repex.py (20260523-0901Z) — repository export for LLMs and humans.

Single-file Python script. Exports a local folder (or a remote repo via
--remote) to docx / xlsx / md / json / odt / ods.

Quick:
    py repex.py . -s agent -o map.md     # LLM-agent orientation, no inline content
    py repex.py . -s llm -o context.md   # paste-into-chat LLM context, full inline
    py repex.py . -o report.docx         # human-readable Word report

Presets (-s / --sections):
    default / all   everything
    agent           LLM with file tools — map only, no inline content
    llm             LLM without tools — map + full content inline
    human           Human reader skim

Format is inferred from the -o / --output extension; -f / --format overrides.

Git discovery: repex looks for .git at the path, at any ancestor, then in
immediate child folders (workspace mode). Files outside every discovered
repo are tagged 'N' (no-repo) instead of T (tracked) / U (untracked).

Full reference: README.md or `py repex.py -h`.
"""

from __future__ import annotations

import argparse
import fnmatch
import hashlib
import json
import os
import re
import subprocess
import sys
import tempfile
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Set, Tuple


__version__ = "20260523-0901Z"


def generator_name() -> str:
    """Author/creator string stamped into output document metadata."""
    return f"repex.py ({__version__})"


DEFAULT_EXTENSIONS = {
    ".py", ".pyw",
    ".cpp", ".cc", ".cxx", ".c", ".h", ".hpp", ".hh",
    ".cs",
    ".java",
    ".js", ".jsx", ".ts", ".tsx",
    ".rs",
    ".go",
    ".rb",
    ".php",
    ".kt", ".kts",
    ".scala",
    ".swift",
    ".html", ".htm", ".css", ".scss", ".sass",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".xml",
    ".md", ".txt", ".rst",
    ".cmake",
    ".sh", ".bat", ".ps1",
    ".sql",
    ".r", ".rmd",
}

DEFAULT_EXCLUDE_DIRS = {
    ".git", ".hg", ".svn",
    "node_modules",
    ".venv", "venv", "__pycache__", ".mypy_cache", ".pytest_cache",
    ".idea", ".vs", ".vscode",
    ".claude",
}

DEFAULT_EXCLUDE_PATTERNS = {
    "*.zip", "*.7z", "*.rar", "*.tar", "*.gz",
    "*.pyc", "*.pyo",
}

SPECIAL_TEXT_FILENAMES = {
    ".gitignore", ".gitattributes", ".editorconfig"
}

TEXT_CANDIDATE_EXTENSIONS = DEFAULT_EXTENSIONS.copy()

TEXT_NAME_WHITELIST = {
    "LICENSE", "LICENCE", "COPYING", "NOTICE", "README", "MAKEFILE", "DOCKERFILE",
    "CHANGELOG", "AUTHORS", "CONTRIBUTORS", "INSTALL",
}

LANGUAGE_KIND_BY_EXT = {
    # programming source code
    ".py": "code", ".pyw": "code",
    ".cpp": "code", ".cc": "code", ".cxx": "code",
    ".c": "code",
    ".h": "code", ".hpp": "code", ".hh": "code",
    ".cs": "code", ".java": "code",
    ".js": "code", ".jsx": "code", ".ts": "code", ".tsx": "code",
    ".rs": "code", ".go": "code", ".rb": "code", ".php": "code",
    ".swift": "code", ".kt": "code", ".kts": "code", ".scala": "code",
    ".r": "code", ".rmd": "docs",
    ".html": "code", ".htm": "code",
    ".css": "code", ".scss": "code", ".sass": "code",
    # human documentation
    ".md": "docs", ".rst": "docs", ".txt": "docs",
    # structured data
    ".json": "data", ".yaml": "data", ".yml": "data",
    ".xml": "data", ".toml": "data",
    ".csv": "data", ".tsv": "data",
    # build / config / scripts
    ".ini": "config", ".cfg": "config", ".conf": "config",
    ".cmake": "config",
    ".sh": "config", ".bat": "config", ".ps1": "config",
    ".sql": "config",
}

# Universal directory-bucket → role mapping. Project-specific overrides should
# be derived from path/content heuristics, not from this table.
ROLE_HINTS = {
    "(repo root)": "other",
    "tests": "test", "test": "test", "spec": "test", "specs": "test",
    "docs": "docs", "doc": "docs",
    "data": "asset", "assets": "asset", "resources": "asset", "resource": "asset",
    "tools": "build", "scripts": "build", "build": "build",
    "packaging": "build", "release": "asset",
    ".github": "config", ".github/workflows": "build",
    "src": "code", "lib": "code", "app": "code", "pkg": "code", "internal": "code",
}

# Per-language function-definition regexes (capturing the function name) used
# both for honest function counts and for harvesting project-local symbols.
# Languages without a pattern get an unknown count (not displayed).
FUNCTION_DEF_PATTERNS_BY_SUFFIX: Dict[str, List["re.Pattern[str]"]] = {
    ".py":   [re.compile(r"^[ \t]*(?:async\s+)?def\s+([A-Za-z_]\w*)\s*\(", re.MULTILINE)],
    ".rs":   [re.compile(r"\bfn\s+([A-Za-z_]\w*)\s*[<(]")],
    ".go":   [re.compile(r"^func\s+(?:\([^)]+\)\s+)?([A-Za-z_]\w*)\s*\(", re.MULTILINE)],
    ".java": [re.compile(
        r"^[ \t]*(?:public|private|protected|static|final|abstract|synchronized|\s)+"
        r"[\w<>\[\],\s]+?\s+([A-Za-z_]\w*)\s*\([^)]*\)\s*(?:throws[^{]+)?\{",
        re.MULTILINE,
    )],
    ".cs":   [re.compile(
        r"^[ \t]*(?:public|private|protected|internal|static|async|virtual|override|"
        r"sealed|abstract|partial|extern|\s)+[\w<>\[\],\s\?]+?\s+([A-Za-z_]\w*)"
        r"\s*\([^)]*\)\s*\{",
        re.MULTILINE,
    )],
    ".js":   [re.compile(r"\bfunction\s+([A-Za-z_]\w*)\s*\("),
              re.compile(
                  r"^[ \t]*(?:export\s+)?(?:const|let|var)\s+([A-Za-z_]\w*)\s*=\s*"
                  r"(?:async\s+)?\([^)]*\)\s*=>",
                  re.MULTILINE,
              )],
    ".cpp":  [re.compile(
        r"\b([A-Za-z_]\w*(?:::[A-Za-z_]\w*)+)\s*\([^)]*\)\s*"
        r"(?:const\s*)?(?:noexcept[^{]*)?\{"
    ),
              re.compile(
                  r"^(?:[A-Za-z_][\w<>\*&\s:]*?\s+)([A-Za-z_]\w*)\s*\([^)]*\)\s*"
                  r"(?:const\s*)?\{",
                  re.MULTILINE,
              )],
    ".c":    [re.compile(
        r"^(?:[A-Za-z_][\w\s\*&]*\s+)([A-Za-z_]\w*)\s*\([^)]*\)\s*\{",
        re.MULTILINE,
    )],
    ".rb":   [re.compile(r"^[ \t]*def\s+(?:self\.)?([A-Za-z_]\w*[\?\!]?)", re.MULTILINE)],
    ".php":  [re.compile(r"\bfunction\s+([A-Za-z_]\w*)\s*\(")],
    ".kt":   [re.compile(r"\bfun\s+(?:[A-Za-z_]\w*\.)?([A-Za-z_]\w*)\s*\(")],
    ".scala":[re.compile(r"\bdef\s+([A-Za-z_]\w*)\s*[\[(:]")],
    ".swift":[re.compile(r"\bfunc\s+([A-Za-z_]\w*)\s*[<(]")],
    ".sh":   [re.compile(r"^[ \t]*([A-Za-z_]\w*)\s*\(\s*\)\s*\{", re.MULTILINE),
              re.compile(r"^[ \t]*function\s+([A-Za-z_]\w*)", re.MULTILINE)],
    ".r":    [re.compile(r"^[ \t]*([A-Za-z_\.][\w\.]*)\s*(?:<-|=)\s*function\s*\(", re.MULTILINE)],
}
for _src_ext, _aliases in [
    (".py",  [".pyw"]),
    (".cpp", [".cc", ".cxx", ".hpp", ".hh", ".h"]),
    (".js",  [".jsx", ".ts", ".tsx"]),
    (".kt",  [".kts"]),
]:
    for _alias in _aliases:
        FUNCTION_DEF_PATTERNS_BY_SUFFIX.setdefault(_alias, FUNCTION_DEF_PATTERNS_BY_SUFFIX[_src_ext])

# Captures any identifier (optionally qualified with `::`) immediately before `(`.
CALL_REGEX = re.compile(r"\b([A-Za-z_]\w*(?:::[A-Za-z_]\w*)*)\s*\(")

# Tokens that look like calls but are not — control flow, declarators, casts.
CONTROL_FLOW_NAMES = {
    "if", "for", "while", "switch", "return", "do", "else", "elif", "when",
    "try", "catch", "finally", "with", "raise", "throw", "yield", "match",
    "case", "func", "fn", "def", "class", "struct", "namespace", "typedef",
    "using", "import", "from", "as", "in", "is", "and", "or", "not", "new",
    "delete", "sizeof", "typeof", "instanceof", "void", "auto", "const",
    "static", "extern", "public", "private", "protected", "internal",
    "lambda", "let", "var", "type", "interface", "enum",
    "assert", "print", "println", "printf", "sprintf", "fprintf",
    "true", "false", "nil", "null", "None", "True", "False",
}

# (regex, label, confidence, applicable file suffixes). Suffix-filtered to avoid
# C++ patterns matching Python files etc.
ENTRY_POINT_PATTERNS = [
    (re.compile(r"\bint\s+(?:WINAPI\s+)?(?:w?WinMain|main)\s*\("),
     "C/C++ main()", "high", {".c", ".cc", ".cpp", ".cxx"}),
    (re.compile(r"^\s*if\s+__name__\s*==\s*['\"]__main__['\"]\s*:", re.MULTILINE),
     "Python __main__ guard", "high", {".py", ".pyw"}),
    (re.compile(r"^def\s+main\s*\(", re.MULTILINE),
     "Python main() function", "medium", {".py", ".pyw"}),
    (re.compile(r"\bfn\s+main\s*\("),
     "Rust main()", "high", {".rs"}),
    (re.compile(r"\bfunc\s+main\s*\("),
     "Go main()", "high", {".go"}),
    (re.compile(r"\bpublic\s+static\s+void\s+main\s*\("),
     "Java main()", "high", {".java"}),
    (re.compile(r"\bstatic\s+(?:async\s+)?(?:void|int|Task)\s+Main\s*\("),
     "C# Main()", "high", {".cs"}),
    (re.compile(r"^export\s+default\b", re.MULTILINE),
     "ES module default export", "low", {".js", ".jsx", ".ts", ".tsx"}),
    (re.compile(r"\bmodule\.exports\s*="),
     "Node.js module export", "low", {".js"}),
    (re.compile(r"^#!.+\b(?:bash|sh|zsh)\b", re.MULTILINE),
     "shell script entry", "medium", {".sh"}),
    (re.compile(r"\b(?:unittest\.main|pytest\.main)\s*\("),
     "Python test runner", "high", {".py"}),
    (re.compile(r"\bcommandArgs\s*\(\s*trailingOnly", re.MULTILINE),
     "R script entry (commandArgs)", "medium", {".r"}),
    (re.compile(r"^\s*main\s*<-\s*function\s*\(", re.MULTILINE),
     "R main() function", "medium", {".r"}),
]


# Document sections, used by the section-aware formats (docx, md, odt).
# Order in ALL_SECTIONS is the canonical render order.
ALL_SECTIONS: Tuple[str, ...] = (
    "glance", "recent", "architecture",
    "entry_points", "trace", "core", "toc", "entries",
)
DEFAULT_SECTIONS: Tuple[str, ...] = ALL_SECTIONS

# Section presets are mnemonic shortcuts for who/what the export is meant for:
#   default / all : everything (handy for hand-off snapshots).
#   llm           : LLM without tools (paste-into-chat). Includes 'entries'
#                   so the model has full file contents in one shot.
#   agent         : LLM with file-access tools (Claude Code, Cursor, etc.).
#                   Skips 'entries' since the agent can Read on demand;
#                   keeps the cross-file intelligence (call sketch, used_by,
#                   entry points) that would otherwise take many greps.
#   human         : Human reader skimming the repo. Overview only, no trace.
SECTION_PRESETS: Dict[str, Tuple[str, ...]] = {
    "default": DEFAULT_SECTIONS,
    "all": ALL_SECTIONS,
    "llm": ("glance", "architecture", "entry_points", "trace", "core", "entries"),
    "agent": ("glance", "architecture", "entry_points", "trace", "core", "toc"),
    "human": ("glance", "recent", "architecture", "toc"),
}

# Output formats. xlsx/json/ods always include every record (no --sections).
SUPPORTED_FORMATS: Tuple[str, ...] = ("docx", "xlsx", "md", "json", "odt", "ods")


def resolve_format(explicit_format: Optional[str], output_path: Optional[str]) -> str:
    """Pick output format: explicit --format wins, else infer from --output
    extension, else default to docx."""
    if explicit_format:
        return explicit_format
    if output_path:
        ext = Path(output_path).suffix.lower().lstrip(".")
        if ext in SUPPORTED_FORMATS:
            return ext
    return "docx"


def resolve_sections(spec: str) -> Set[str]:
    """Parse --sections value. Accepts a preset name ('default', 'all', 'llm',
    'human'), explicit csv of section names, or additive/subtractive csv mixing
    presets with '+name' / '-name'."""
    spec = (spec or "").strip()
    if not spec:
        return set(DEFAULT_SECTIONS)

    chosen: Optional[Set[str]] = None
    for raw in spec.split(","):
        token = raw.strip()
        if not token:
            continue
        if token in SECTION_PRESETS:
            preset = set(SECTION_PRESETS[token])
            chosen = preset if chosen is None else chosen | preset
            continue
        sign = ""
        name = token
        if token[0] in "+-":
            sign = token[0]
            name = token[1:].strip()
        if name not in ALL_SECTIONS:
            raise ValueError(
                f"unknown section: {token!r}. "
                f"Valid presets: {', '.join(SECTION_PRESETS)}. "
                f"Valid names: {', '.join(ALL_SECTIONS)}"
            )
        if chosen is None:
            chosen = set()
        if sign == "-":
            chosen.discard(name)
        else:
            chosen.add(name)
    return chosen if chosen is not None else set(DEFAULT_SECTIONS)


# ---------------- core helpers ----------------

# Mapping of XML-incompatible codepoints -> '?'. XML 1.0 allows:
#   #x9, #xA, #xD, #x20-#xD7FF, #xE000-#xFFFD, #x10000-#x10FFFF
# Everything else (control chars, surrogates, 0xFFFE/0xFFFF) must be scrubbed.
# Using str.translate runs the substitution in C, ~100x faster than the
# per-character Python loop for large embedded file contents.
def _build_xml_scrub_table() -> Dict[int, int]:
    table: Dict[int, int] = {}
    bad: List[int] = []
    # C0 controls except TAB/LF/CR.
    for code in range(0x00, 0x20):
        if code not in (0x09, 0x0A, 0x0D):
            bad.append(code)
    # Surrogate range.
    bad.extend(range(0xD800, 0xE000))
    # Non-characters at the BMP end.
    bad.append(0xFFFE)
    bad.append(0xFFFF)
    q = ord("?")
    for code in bad:
        table[code] = q
    return table


_XML_SCRUB_TABLE: Dict[int, int] = _build_xml_scrub_table()


def sanitize_xml_compatible_text(text: str) -> str:
    """Replace XML-incompatible control / surrogate / non-character codepoints
    with '?', preserving tabs / newlines / CR. C-speed via str.translate."""
    if not text:
        return text
    return text.translate(_XML_SCRUB_TABLE)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Export a local repository to docx, xlsx, md, json, odt, or ods "
            "for LLM context ingestion or human overview."
        )
    )
    parser.add_argument(
        "repo",
        nargs="?",
        default=None,
        help=(
            "Path to the repository/folder. Optional when --remote is set "
            "(in which case the remote is cloned into a tempdir)."
        ),
    )
    parser.add_argument(
        "--remote",
        default=None,
        help=(
            "Clone a remote repository into a tempdir and export it. "
            "Accepts 'owner/repo' (GitHub shorthand) or any clone URL. "
            "The tempdir is removed when the export finishes."
        ),
    )
    parser.add_argument(
        "-f", "--format",
        choices=["docx", "xlsx", "md", "json", "odt", "ods"],
        default=None,
        help=(
            "Output format. If omitted, inferred from the -o / --output "
            "extension. If neither is set, defaults to docx."
        ),
    )
    parser.add_argument(
        "-o", "--output",
        help=(
            "Output file path. Default: <repo_name>_code_export.<format>. "
            "If -f / --format is omitted, this extension decides the format."
        ),
    )
    parser.add_argument(
        "--ext",
        nargs="*",
        default=None,
        help="Only include these extensions, e.g. --ext .py .cpp .h .md",
    )
    parser.add_argument(
        "--exclude-dir",
        nargs="*",
        default=[],
        help="Additional directory names to exclude.",
    )
    parser.add_argument(
        "--exclude-pattern",
        nargs="*",
        default=[],
        help="Additional filename patterns to exclude, e.g. *.min.js *.lock",
    )
    parser.add_argument(
        "--max-text-file-kb",
        type=float,
        default=512.0,
        help=(
            "For text files larger than this size in kB, do not include full content; "
            "list metadata only. Default: 512"
        ),
    )
    parser.add_argument(
        "--encoding",
        default="utf-8",
        help="Preferred text encoding. Default: utf-8",
    )
    parser.add_argument(
        "--include-no-extension",
        action="store_true",
        help="Also consider files without extension as possible text files.",
    )
    parser.add_argument(
        "--no-gitignore",
        action="store_true",
        help=(
            "Disable .gitignore filtering (default: enabled when the target "
            "is a git repository). Use this to include files your .gitignore "
            "would normally hide."
        ),
    )
    parser.add_argument(
        "--since",
        default=None,
        help=(
            "Restrict the export to files changed since <ref> (any git "
            "revision: branch, tag, commit, e.g. 'main', 'HEAD~5', 'v1.2'). "
            "Includes committed diff vs the ref, working-tree changes, and "
            "untracked-not-ignored files. Requires a git repository."
        ),
    )
    parser.add_argument(
        "--strip-comments",
        action="store_true",
        help=(
            "Strip line and block comments from code files before embedding "
            "them in the export. Typically saves 15-30%% tokens on the "
            "'llm' preset. Strings are preserved."
        ),
    )
    parser.add_argument(
        "--token-budget",
        type=int,
        default=None,
        help=(
            "Markdown/JSON only. Target token count for the final output. "
            "If the rendered export exceeds the budget, content is dropped "
            "from the lowest-ranked files (by used_by and size) until the "
            "budget is met. Without tiktoken installed, a 4-char/token "
            "estimate is used."
        ),
    )
    parser.add_argument(
        "--token-model",
        default="gpt-4o",
        help=(
            "Model name passed to tiktoken for token counting. Default: "
            "gpt-4o. The cl100k tokenizer is a reasonable proxy for Claude."
        ),
    )
    parser.add_argument(
        "--clipboard",
        action="store_true",
        help=(
            "Markdown/JSON only. Also copy the rendered output to the system "
            "clipboard. Uses pyperclip if available, else native commands "
            "(clip.exe / pbcopy / wl-copy / xclip)."
        ),
    )
    parser.add_argument(
        "-s", "--sections",
        default="default",
        help=(
            "Section selection for docx/md/odt (xlsx/json/ods always include "
            "everything). Comma-separated. "
            "Presets: " + ", ".join(SECTION_PRESETS) + ". "
            "Names: " + ", ".join(ALL_SECTIONS) + ". "
            "Use '+name' / '-name' to add/remove from a preset "
            "(e.g. 'llm,+toc' or 'all,-entries'). "
            "Default: default"
        ),
    )
    return parser.parse_args()


def is_git_repo(repo: Path) -> bool:
    """True iff `repo` itself is a git work tree (i.e. has a `.git` entry)."""
    return (repo / ".git").exists()


# Kind labels for discover_git_roots: where the .git directory sits
# relative to the user-supplied repo path.
GIT_ROOT_KIND_SELF = "self"          # repo itself has .git
GIT_ROOT_KIND_ANCESTOR = "ancestor"  # found by walking up from repo
GIT_ROOT_KIND_CHILD = "child"        # immediate-child subfolder of repo


def discover_git_roots(repo: Path) -> List[Tuple[Path, str]]:
    """Find every git work tree that touches the user-supplied `repo` path.

    Search order:
      1. `repo` itself. If it has .git, stop — that's the only relevant repo.
      2. Otherwise walk UP the parents; the first ancestor with .git is used.
      3. Otherwise scan IMMEDIATE children of `repo` for subfolders with .git.

    Returns a list of (root_path, kind) tuples. Empty when nothing is git-
    backed. Walk-up handles `repex C:/proj/src` when the repo is at C:/proj.
    Walk-down handles workspaces where C:/work contains C:/work/projA/.git
    and C:/work/projB/.git but C:/work itself isn't a repo.
    """
    if (repo / ".git").exists():
        return [(repo, GIT_ROOT_KIND_SELF)]
    # Walk up.
    cur = repo.parent
    while cur != cur.parent:
        if (cur / ".git").exists():
            return [(cur, GIT_ROOT_KIND_ANCESTOR)]
        cur = cur.parent
    # Walk down (immediate children only).
    roots: List[Tuple[Path, str]] = []
    try:
        for child in sorted(repo.iterdir(), key=lambda p: p.name.lower()):
            if child.is_dir() and (child / ".git").exists():
                roots.append((child, GIT_ROOT_KIND_CHILD))
    except (OSError, PermissionError):
        pass
    return roots


def assign_git_origin(
    file_path: Path,
    roots: Sequence[Tuple[Path, str]],
) -> Tuple[Optional[Path], str]:
    """Pick the git root that contains `file_path`, if any.
    Returns (root_path or None, rel_path_inside_root or '')."""
    for root, _ in roots:
        try:
            rel = file_path.relative_to(root).as_posix()
            return root, rel
        except ValueError:
            continue
    return None, ""


def render_repo_label(root: Path, repo: Path) -> str:
    """Short, display-friendly label for a git root relative to `repo`.
    Used in TOC markers and summary lines."""
    try:
        rel = root.relative_to(repo).as_posix()
        return rel if rel and rel != "." else root.name
    except ValueError:
        # Ancestor: not under repo. Show its basename.
        return root.name


def render_toc_marker(
    record: Dict[str, object],
    repo: Path,
    multi_root: bool,
) -> str:
    """One-letter (or 'T:label' / 'U:label') tag describing the record's
    relationship to git: T=tracked, U=untracked, N=no-repo (file outside
    any discovered git tree). When `multi_root` is True, T/U also carry
    a repo label so the reader can distinguish workspaces."""
    origin = record.get("git_origin")
    if origin is None:
        return "N"
    base = "T" if record.get("git_tracked") else "U"
    if multi_root:
        return f"{base}:{render_repo_label(origin, repo)}"
    return base


def format_git_roots_summary(
    roots: Sequence[Tuple[Path, str]],
    repo: Path,
    tracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    untracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    local_only_records: Sequence[Tuple[Path, Dict[str, object]]],
) -> str:
    """One-line description of the git topology for the run. Embedded in
    every export's header so a reader can tell whether they're looking at
    a single repo, a workspace of subrepos, or a no-repo folder."""
    if not roots:
        return "Git: no repository detected (local-only mode)."
    parts: List[str] = []
    for root, kind in roots:
        label = render_repo_label(root, repo)
        t = sum(1 for _, r in tracked_records if r.get("git_origin") == root)
        u = sum(1 for _, r in untracked_records if r.get("git_origin") == root)
        parts.append(f"{label} ({kind}; T:{t}, U:{u})")
    no_repo = len(local_only_records)
    tail = f"; no-repo files: {no_repo}" if no_repo > 0 else ""
    return "Git roots: " + "; ".join(parts) + tail


def run_git_command(repo: Path, args: List[str]) -> str:
    try:
        result = subprocess.run(
            ["git", "-C", str(repo)] + args,
            capture_output=True,
            text=True,
            check=True,
        )
        return result.stdout.strip()
    except FileNotFoundError as exc:
        raise RuntimeError("Git is not installed or not available in PATH.") from exc
    except subprocess.CalledProcessError as exc:
        msg = exc.stderr.strip() or str(exc)
        raise RuntimeError(f"`git {' '.join(args)}` failed: {msg}") from exc


def get_git_tracked_relpaths(repo: Path) -> Set[Path]:
    output = run_git_command(repo, ["ls-files"])
    relpaths: Set[Path] = set()
    for line in output.splitlines():
        line = line.strip()
        if line:
            relpaths.add(Path(line))
    return relpaths


def get_git_branch_name(repo: Path) -> str:
    try:
        branch = run_git_command(repo, ["rev-parse", "--abbrev-ref", "HEAD"])
        return branch if branch else "(unknown)"
    except RuntimeError:
        return "(unknown)"


def get_git_head_commit(repo: Path) -> str:
    try:
        commit = run_git_command(repo, ["rev-parse", "--short", "HEAD"])
        return commit if commit else "(unknown)"
    except RuntimeError:
        return "(unknown)"


def get_git_head_commit_datetime_utc(repo: Path) -> str:
    try:
        dt = run_git_command(repo, ["log", "-1", "--format=%cd", "--date=format:%Y-%m-%d %H:%M:%S UTC"])
        return dt if dt else "(unknown)"
    except RuntimeError:
        return "(unknown)"


def should_exclude_by_pattern(path: Path, patterns: Sequence[str]) -> bool:
    name = path.name
    rel = path.as_posix()
    for pattern in patterns:
        if fnmatch.fnmatch(name, pattern) or fnmatch.fnmatch(rel, pattern):
            return True
    return False


def _filter_by_gitignore_single(root: Path, paths: Sequence[Path]) -> List[Path]:
    """Apply one git work tree's .gitignore to the subset of `paths` that
    live under `root`. Paths outside `root` are passed through untouched.
    Used internally by filter_by_gitignore_for_roots."""
    if not paths:
        return list(paths)
    under_root: List[Path] = []
    outside: List[Path] = []
    for p in paths:
        try:
            p.relative_to(root)
            under_root.append(p)
        except ValueError:
            outside.append(p)
    if not under_root:
        return list(paths)
    try:
        rel_lines = "\n".join(p.relative_to(root).as_posix() for p in under_root)
        proc = subprocess.run(
            ["git", "-C", str(root), "check-ignore", "--stdin"],
            input=rel_lines,
            capture_output=True,
            text=True,
            check=False,
        )
        if proc.returncode > 1:
            # Real git error — preserve everything rather than fail loudly.
            return list(paths)
        ignored = {line.strip() for line in proc.stdout.splitlines() if line.strip()}
    except FileNotFoundError:
        return list(paths)
    if not ignored:
        return list(paths)
    kept_under_root = [
        p for p in under_root if p.relative_to(root).as_posix() not in ignored
    ]
    # Preserve original order.
    kept_set = set(kept_under_root)
    return [p for p in paths if p in kept_set or p in outside]


def filter_by_gitignore_for_roots(
    roots: Sequence[Tuple[Path, str]],
    paths: Sequence[Path],
) -> List[Path]:
    """Drop paths that any discovered git root's .gitignore would exclude.
    A path inside multiple roots (shouldn't happen with our discovery, but
    just in case) survives only if every relevant root keeps it."""
    result = list(paths)
    for root, _ in roots:
        result = _filter_by_gitignore_single(root, result)
    return result


def get_changed_paths_for_roots(
    roots: Sequence[Tuple[Path, str]],
    ref: str,
) -> Set[Path]:
    """For each git root, collect the files changed since <ref>: committed
    diff (ref..HEAD) + working-tree diff vs HEAD + untracked-not-ignored.
    Returns absolute Paths so callers can match against the file list
    without worrying about which root a path lives in."""
    if not roots:
        raise RuntimeError("--since requires at least one git repository.")
    changed: Set[Path] = set()
    for root, _ in roots:
        try:
            out = run_git_command(root, ["diff", "--name-only", f"{ref}..HEAD"])
            for line in out.splitlines():
                line = line.strip()
                if line:
                    changed.add(root / line)
        except RuntimeError as exc:
            raise RuntimeError(
                f"--since: cannot resolve ref {ref!r} in {root}: {exc}"
            ) from exc
        try:
            out = run_git_command(root, ["diff", "--name-only", "HEAD"])
            for line in out.splitlines():
                line = line.strip()
                if line:
                    changed.add(root / line)
        except RuntimeError:
            pass
        try:
            out = run_git_command(
                root, ["ls-files", "--others", "--exclude-standard"]
            )
            for line in out.splitlines():
                line = line.strip()
                if line:
                    changed.add(root / line)
        except RuntimeError:
            pass
    return changed


def clone_remote_to_tempdir(remote: str) -> Tuple[Path, "tempfile.TemporaryDirectory"]:
    """Shallow-clone a remote into a tempdir. The TemporaryDirectory must be
    kept alive (and explicitly cleaned up) by the caller. Accepts either a
    full URL or 'owner/repo' shorthand (assumed to be a public github.com
    repository)."""
    if "/" in remote and "://" not in remote and not remote.startswith("git@"):
        url = f"https://github.com/{remote}.git"
    else:
        url = remote
    tmpdir = tempfile.TemporaryDirectory(prefix="repex-remote-")
    target = Path(tmpdir.name) / "repo"
    try:
        subprocess.run(
            ["git", "clone", "--depth", "1", url, str(target)],
            check=True,
            capture_output=True,
            text=True,
        )
    except FileNotFoundError as exc:
        tmpdir.cleanup()
        raise RuntimeError("Git is not installed or not available in PATH.") from exc
    except subprocess.CalledProcessError as exc:
        tmpdir.cleanup()
        msg = exc.stderr.strip() or str(exc)
        raise RuntimeError(f"Cloning {url!r} failed: {msg}") from exc
    return target, tmpdir


def copy_text_to_clipboard(text: str) -> bool:
    """Try to copy text to the system clipboard. Returns True on success.
    Prefers pyperclip if installed; falls back to platform native commands
    (clip.exe on Windows, pbcopy on macOS, xclip/wl-copy on Linux)."""
    try:
        import pyperclip  # type: ignore
        pyperclip.copy(text)
        return True
    except Exception:
        pass
    # Platform-native fallbacks.
    candidates: List[List[str]]
    if sys.platform.startswith("win"):
        candidates = [["clip"]]
    elif sys.platform == "darwin":
        candidates = [["pbcopy"]]
    else:
        candidates = [["wl-copy"], ["xclip", "-selection", "clipboard"]]
    for cmd in candidates:
        try:
            proc = subprocess.run(cmd, input=text, text=True, check=False)
            if proc.returncode == 0:
                return True
        except FileNotFoundError:
            continue
    return False


# Per-language comment-stripper regexes. Strips comments only; preserves
# strings (unlike _strip_for_call_detection which is destructive — it
# removes strings too so call-graph regex doesn't fire on identifiers
# inside literals or comments).
# Patterns are applied in order; multiline pattern (block comments) first
# avoids the line-comment pattern eating contents of unterminated blocks.
_LINE_COMMENT_HASH = {
    ".py", ".pyw", ".rb", ".sh", ".r", ".ps1",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".cmake",
}
_LINE_COMMENT_DOUBLE_SLASH = {
    ".c", ".cpp", ".cc", ".cxx", ".h", ".hpp", ".hh",
    ".cs", ".java", ".js", ".jsx", ".ts", ".tsx",
    ".go", ".rs", ".kt", ".kts", ".scala", ".swift",
    ".php", ".css", ".scss", ".sass",
}
_BLOCK_COMMENT_C_FAMILY = _LINE_COMMENT_DOUBLE_SLASH | {".php"}
_HTML_XML_COMMENT = {".html", ".htm", ".xml"}
_SQL_DASH_COMMENT = {".sql"}


def strip_comments_only(text: str, suffix: str) -> str:
    """Remove comments from `text` based on file `suffix`. Leaves strings
    intact. Used by --strip-comments to shrink LLM contexts ~15-30% without
    altering executable semantics.

    Conservative: when in doubt, leaves the line alone. We do not parse
    strings, so a '#' inside a Python string would be wrongly treated as a
    comment-start; in practice the saving is high and the breakage low for
    text that an LLM reads (not executes)."""
    suffix = suffix.lower()
    if not text:
        return text
    out = text
    # Block comments first.
    if suffix in _BLOCK_COMMENT_C_FAMILY:
        out = re.sub(r"/\*.*?\*/", "", out, flags=re.DOTALL)
    if suffix in _HTML_XML_COMMENT:
        out = re.sub(r"<!--.*?-->", "", out, flags=re.DOTALL)
    # Line comments (drop the comment, keep the newline so line numbers
    # in tracebacks/links stay roughly aligned).
    if suffix in _LINE_COMMENT_HASH:
        out = re.sub(r"(^|\s)#.*?$", r"\1", out, flags=re.MULTILINE)
    if suffix in _LINE_COMMENT_DOUBLE_SLASH:
        out = re.sub(r"//.*?$", "", out, flags=re.MULTILINE)
    if suffix in _SQL_DASH_COMMENT:
        out = re.sub(r"--.*?$", "", out, flags=re.MULTILINE)
    # Collapse any runs of blank lines created by removal.
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out


# tiktoken encoder cache. The library itself caches by name internally,
# but Python's import machinery and the per-call dict lookup still add up
# when annotating thousands of files. We resolve once per model and reuse.
# _TIKTOKEN_PROBE: None = not yet tried, False = import failed, True = ok.
_TIKTOKEN_PROBE: Optional[bool] = None
_TOKEN_ENCODERS: Dict[str, object] = {}


def _get_token_encoder(model_hint: str):
    """Lazily resolve and cache a tiktoken encoder for `model_hint`. Returns
    None when tiktoken isn't installed (probed once) or the model is unknown
    even with the cl100k fallback."""
    global _TIKTOKEN_PROBE
    if _TIKTOKEN_PROBE is False:
        return None
    cached = _TOKEN_ENCODERS.get(model_hint)
    if cached is not None:
        return cached
    try:
        import tiktoken  # type: ignore
    except ImportError:
        _TIKTOKEN_PROBE = False
        return None
    _TIKTOKEN_PROBE = True
    try:
        try:
            enc = tiktoken.encoding_for_model(model_hint)
        except KeyError:
            enc = tiktoken.get_encoding("cl100k_base")
    except Exception:
        return None
    _TOKEN_ENCODERS[model_hint] = enc
    return enc


def count_tokens(text: str, model_hint: str = "gpt-4o") -> Optional[int]:
    """Estimate token count using tiktoken. Returns None if tiktoken is
    unavailable so callers can warn instead of crashing. The exact tokenizer
    differs across model families (OpenAI cl100k for gpt-4o; Claude uses its
    own); cl100k is close enough for budgeting purposes."""
    enc = _get_token_encoder(model_hint)
    if enc is None:
        return None
    try:
        return len(enc.encode(text))
    except Exception:
        return None


def estimate_tokens_rough(text: str) -> int:
    """Rough offline token estimate (~4 chars/token). Used when tiktoken
    is not installed. Good enough for budget pruning where the goal is
    'roughly fit' rather than exact accounting."""
    return max(1, len(text) // 4)


def rank_records_for_pruning(
    records: Sequence[Tuple[Path, Dict[str, object]]],
) -> List[Tuple[Path, Dict[str, object]]]:
    """Return records sorted highest-value first.
    Score = used_by count, tie-broken by smaller size (smaller wins; we
    keep cheap-but-referenced files first when the budget is tight)."""
    def key(item: Tuple[Path, Dict[str, object]]) -> Tuple[int, int]:
        _path, record = item
        used_by = record.get("used_by") or []
        return (-len(list(used_by)), int(record.get("size_bytes", 0)))
    return sorted(records, key=key)


def looks_like_text(path: Path, sample_size: int = 4096) -> bool:
    try:
        with path.open("rb") as f:
            chunk = f.read(sample_size)
        if b"\x00" in chunk:
            return False
        return True
    except Exception:
        return False


def is_text_file(path: Path, include_no_extension: bool = False) -> bool:
    suffix = path.suffix.lower()

    if path.name in SPECIAL_TEXT_FILENAMES:
        return True

    if suffix in TEXT_CANDIDATE_EXTENSIONS:
        return True

    if path.name.upper() in TEXT_NAME_WHITELIST:
        return True

    if suffix:
        return looks_like_text(path)

    if include_no_extension:
        return looks_like_text(path)

    return False


def collect_all_files(
    repo: Path,
    allowed_extensions: Optional[Set[str]],
    exclude_dirs: Set[str],
    exclude_patterns: Set[str],
) -> List[Path]:
    candidates: List[Path] = []

    for root, dirs, files in os.walk(repo):
        dirs[:] = [d for d in dirs if d not in exclude_dirs]
        for file in files:
            candidates.append(Path(root) / file)

    selected: List[Path] = []
    for path in candidates:
        if not path.is_file():
            continue

        rel = path.relative_to(repo)

        if any(part in exclude_dirs for part in rel.parts[:-1]):
            continue

        if should_exclude_by_pattern(rel, exclude_patterns):
            continue

        if allowed_extensions is not None:
            suffix = path.suffix.lower()
            if suffix:
                if suffix not in allowed_extensions:
                    continue
            else:
                continue

        selected.append(path)

    selected.sort(key=lambda p: p.relative_to(repo).as_posix().lower())
    return selected


def read_text_file(path: Path, preferred_encoding: str) -> str:
    encodings_to_try = [preferred_encoding, "utf-8", "utf-8-sig", "cp1252", "latin-1"]
    seen = set()

    for enc in encodings_to_try:
        enc_norm = enc.lower()
        if enc_norm in seen:
            continue
        seen.add(enc_norm)

        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            return f"[ERROR reading file: {exc}]"

    return "[ERROR reading file: unable to decode text content]"


def count_lines(text: str) -> int:
    if text == "":
        return 0
    return text.count("\n") + 1


def detect_language(path: Path) -> str:
    suffix = path.suffix.lower()
    mapping = {
        ".py": "Python", ".pyw": "Python",
        ".cpp": "C++", ".cc": "C++", ".cxx": "C++", ".hpp": "C++ Header", ".hh": "C++ Header",
        ".c": "C", ".h": "C/C++ Header",
        ".cs": "C#",
        ".java": "Java",
        ".js": "JavaScript", ".jsx": "JavaScript/JSX",
        ".ts": "TypeScript", ".tsx": "TypeScript/TSX",
        ".rs": "Rust",
        ".go": "Go",
        ".rb": "Ruby",
        ".php": "PHP",
        ".kt": "Kotlin", ".kts": "Kotlin Script",
        ".scala": "Scala",
        ".swift": "Swift",
        ".html": "HTML", ".htm": "HTML",
        ".css": "CSS", ".scss": "SCSS", ".sass": "Sass",
        ".json": "JSON", ".yaml": "YAML", ".yml": "YAML", ".xml": "XML", ".toml": "TOML",
        ".md": "Markdown", ".txt": "Text", ".rst": "reStructuredText",
        ".cmake": "CMake",
        ".sh": "Shell", ".bat": "Batch", ".ps1": "PowerShell",
        ".sql": "SQL",
        ".r": "R", ".rmd": "RMarkdown",
        ".ico": "Icon",
        ".exe": "Executable",
        ".dll": "Dynamic Link Library",
        ".rc": "Resource Script",
        ".in": "Template",
        ".example": "Example",
    }

    if path.name in SPECIAL_TEXT_FILENAMES:
        return "Config"

    if suffix in mapping:
        return mapping[suffix]

    if path.name == "CMakeLists.txt":
        return "CMake"

    if path.name.upper() in TEXT_NAME_WHITELIST:
        return "Text"

    if suffix:
        return f"{suffix} file"

    return "No extension"


def classify_language_kind(path: Path) -> str:
    """Bucket a file into 'code' / 'docs' / 'data' / 'config' / 'other'."""
    suffix = path.suffix.lower()
    name = path.name
    name_upper = name.upper()

    if name in SPECIAL_TEXT_FILENAMES:
        return "config"
    if name in {"Dockerfile", "Makefile", "CMakeLists.txt"} or name_upper == "DOCKERFILE":
        return "config"
    if name_upper in TEXT_NAME_WHITELIST:
        return "docs"

    if suffix in LANGUAGE_KIND_BY_EXT:
        return LANGUAGE_KIND_BY_EXT[suffix]
    return "other"


def format_file_size_kb(num_bytes: int) -> str:
    return f"{num_bytes / 1024.0:.2f} kB"


def format_mtime_utc(ts: float) -> str:
    return datetime.fromtimestamp(ts, timezone.utc).strftime("%Y%m%d-%H%MZ")


def describe_non_text_file(path: Path, file_size_bytes: int) -> str:
    return f"[binary file omitted] {path.name} | {format_file_size_kb(file_size_bytes)} | {path.suffix or 'no extension'}"


def describe_large_text_file(path: Path, file_size_bytes: int, max_text_file_kb: float) -> str:
    return (
        f"[text file omitted: too large] {path.name} | {format_file_size_kb(file_size_bytes)} | "
        f"limit={max_text_file_kb:.2f} kB | {path.suffix or 'no extension'}"
    )


def quick_content_hash(path: Path, max_bytes: int = 65536) -> str:
    sha1 = hashlib.sha1()
    try:
        with path.open("rb") as f:
            remaining = max_bytes
            while remaining > 0:
                chunk = f.read(min(8192, remaining))
                if not chunk:
                    break
                sha1.update(chunk)
                remaining -= len(chunk)
        return sha1.hexdigest()[:12]
    except Exception:
        return "(unavailable)"


def infer_directory_bucket(rel_path: str) -> str:
    parts = rel_path.split("/")

    if len(parts) == 1:
        return "(repo root)"

    if parts[0] == "src":
        if len(parts) >= 3 and parts[1] in {"app", "core", "ui", "world", "data", "serial", "input", "platform"}:
            return "/".join(parts[:2])
        return "src"

    if parts[0] == ".github":
        if len(parts) >= 2 and parts[1] == "workflows":
            return ".github/workflows"
        return ".github"

    if parts[0] in {"data", "tests", "tools", "docs", "resources", "packaging", "release"}:
        return parts[0]

    return parts[0]


def infer_role_hint(rel_path: str, text: str = "") -> str:
    """Project-agnostic role classification: test/docs/config/build/asset/code/other."""
    rel = rel_path.replace("\\", "/").lower()
    name = Path(rel).name

    if (
        rel.startswith("tests/") or rel.startswith("test/")
        or "/tests/" in rel or "/test/" in rel
        or name.startswith("test_")
        or name.endswith(("_test.py", "_test.go", "_test.cpp", "_tests.cpp",
                          ".test.js", ".test.jsx", ".test.ts", ".test.tsx",
                          ".spec.js", ".spec.jsx", ".spec.ts", ".spec.tsx"))
    ):
        return "test"

    if (
        rel.startswith("docs/") or rel.startswith("doc/")
        or name in {"readme.md", "readme.rst", "readme.txt", "readme",
                    "changelog.md", "changelog.rst", "changelog",
                    "license", "license.md", "licence", "notice", "notice.md",
                    "authors", "contributors"}
    ):
        return "docs"

    if (
        name in {".gitignore", ".gitattributes", ".editorconfig",
                 ".dockerignore", "dockerfile"}
        or rel.startswith(".github/") and not rel.startswith(".github/workflows/")
    ):
        return "config"

    if (
        name == "cmakelists.txt" or name.endswith(".cmake")
        or name == "makefile"
        or rel.startswith(".github/workflows/")
        or rel.startswith("tools/") or rel.startswith("scripts/")
        or rel.startswith("build/") or rel.startswith("packaging/")
    ):
        return "build"

    if (
        rel.startswith("data/") or rel.startswith("assets/")
        or rel.startswith("resources/") or rel.startswith("resource/")
        or rel.startswith("release/")
    ):
        return "asset"

    if (
        rel.startswith("src/") or rel.startswith("lib/")
        or rel.startswith("app/") or rel.startswith("pkg/")
        or rel.startswith("internal/")
    ):
        return "code"

    bucket = infer_directory_bucket(rel)
    return ROLE_HINTS.get(bucket, "other")


def extract_include_headers(text: str) -> List[str]:
    deps: List[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        m = re.match(r'#include\s+"([^"]+)"', stripped)
        if m:
            deps.append(m.group(1))
    return deps[:24]


def include_path_to_module_name(include_path: str) -> str:
    normalized = include_path.replace("\\", "/").strip()
    name = normalized.split("/")[-1]

    for suffix in (".hpp", ".hh", ".h", ".cpp", ".cc", ".cxx", ".c", ".inl"):
        if name.lower().endswith(suffix):
            name = name[: -len(suffix)]
            break

    if name in {"", ".", ".."}:
        return include_path

    return name


def estimate_complexity_metrics(text: str, suffix: str = "") -> Dict[str, object]:
    lines = count_lines(text)
    class_count = len(re.findall(r"\bclass\s+[A-Za-z_]\w*", text))
    struct_count = len(re.findall(r"\bstruct\s+[A-Za-z_]\w*", text))

    suffix = (suffix or "").lower()
    patterns = FUNCTION_DEF_PATTERNS_BY_SUFFIX.get(suffix)
    if patterns:
        seen_names: Set[str] = set()
        function_count = 0
        for pattern in patterns:
            for match in pattern.finditer(text):
                name = match.group(1)
                key = f"{name}@{match.start()}"
                if key in seen_names:
                    continue
                seen_names.add(key)
                function_count += 1
        functions_known = True
    else:
        function_count = 0
        functions_known = False

    return {
        "lines": lines,
        "classes": class_count,
        "structs": struct_count,
        "functions": function_count,
        "functions_known": functions_known,
    }


def detect_entry_point_signals(text: str, rel_path: str = "") -> List[Tuple[str, str]]:
    suffix = Path(rel_path.replace("\\", "/")).suffix.lower()
    if not suffix:
        return []

    matches: List[Tuple[str, str]] = []
    for pattern, label, confidence, suffixes in ENTRY_POINT_PATTERNS:
        if suffix not in suffixes:
            continue
        if pattern.search(text):
            matches.append((label, confidence))

    deduped: List[Tuple[str, str]] = []
    seen = set()
    for item in matches:
        if item not in seen:
            seen.add(item)
            deduped.append(item)
    return deduped


def get_file_record(
    path: Path,
    rel_path: str,
    preferred_encoding: str,
    max_text_file_kb: float,
    include_no_extension: bool,
) -> Dict[str, object]:
    try:
        stat = path.stat()
        file_size_bytes = stat.st_size
        mtime_utc = format_mtime_utc(stat.st_mtime)
    except OSError:
        file_size_bytes = 0
        mtime_utc = "(unavailable)"

    file_type = detect_language(path)
    text_mode = is_text_file(path, include_no_extension=include_no_extension)
    max_text_file_bytes = int(max_text_file_kb * 1024.0)

    line_count: Optional[int] = None
    omission_reason = ""
    dependencies_raw: List[str] = []
    complexity: Dict[str, object] = {
        "lines": 0, "classes": 0, "structs": 0, "functions": 0, "functions_known": False,
    }
    entry_signals: List[Tuple[str, str]] = []

    if text_mode and file_size_bytes <= max_text_file_bytes:
        content = read_text_file(path, preferred_encoding)
        content_included = True
        line_count = count_lines(content)
        dependencies_raw = extract_include_headers(content)
        complexity = estimate_complexity_metrics(content, path.suffix.lower())
        entry_signals = detect_entry_point_signals(content, rel_path)
    elif text_mode:
        content = describe_large_text_file(path, file_size_bytes, max_text_file_kb)
        content_included = False
        omission_reason = f"text file exceeds {max_text_file_kb:.2f} kB threshold"
    else:
        content = describe_non_text_file(path, file_size_bytes)
        content_included = False
        omission_reason = "binary or non-text file"

    return {
        "type": file_type,
        "kind": classify_language_kind(path),
        "size_bytes": file_size_bytes,
        "size_human": format_file_size_kb(file_size_bytes),
        "mtime_utc": mtime_utc,
        "quick_hash": quick_content_hash(path),
        "is_text": text_mode,
        "line_count": line_count,
        "content_included": content_included,
        "omission_reason": omission_reason,
        "content_or_note": content,
        "role_hint": infer_role_hint(rel_path, content if text_mode and file_size_bytes <= max_text_file_bytes else ""),
        "dependencies_raw": dependencies_raw,
        "dependencies_resolved": [],
        "used_by": [],
        "complexity": complexity,
        "entry_signals": entry_signals,
        # Git provenance — populated by build_records_for_files when given
        # the discovered roots. None / False means the file is outside any
        # repo (workspace-mode "no-repo" bucket).
        "git_origin": None,        # Optional[Path] — root of containing git repo
        "git_tracked": False,      # bool — tracked in that repo
    }


def summarize_records(records: Sequence[Dict[str, object]]) -> Dict[str, object]:
    lines_by_kind: Dict[str, int] = defaultdict(int)
    lines_by_language: Dict[str, int] = defaultdict(int)
    summary: Dict[str, object] = {
        "file_count": 0,
        "text_count": 0,
        "binary_count": 0,
        "content_included_count": 0,
        "content_omitted_count": 0,
        "total_size_bytes": 0,
        "lines_total": 0,
        "lines_by_kind": lines_by_kind,
        "lines_by_language": lines_by_language,
    }
    for record in records:
        summary["file_count"] = int(summary["file_count"]) + 1
        summary["total_size_bytes"] = int(summary["total_size_bytes"]) + int(record["size_bytes"])
        if bool(record["is_text"]):
            summary["text_count"] = int(summary["text_count"]) + 1
        else:
            summary["binary_count"] = int(summary["binary_count"]) + 1
        if bool(record["content_included"]):
            summary["content_included_count"] = int(summary["content_included_count"]) + 1
        else:
            summary["content_omitted_count"] = int(summary["content_omitted_count"]) + 1
        line_count = record.get("line_count")
        if isinstance(line_count, int):
            summary["lines_total"] = int(summary["lines_total"]) + line_count
            kind = str(record.get("kind", "other"))
            lines_by_kind[kind] += line_count
            lines_by_language[str(record.get("type", "Unknown"))] += line_count
    return summary


def build_records_for_files(
    repo: Path,
    files: Sequence[Path],
    preferred_encoding: str,
    max_text_file_kb: float,
    include_no_extension: bool,
    roots: Optional[Sequence[Tuple[Path, str]]] = None,
    tracked_per_root: Optional[Dict[Path, Set[str]]] = None,
) -> List[Tuple[Path, Dict[str, object]]]:
    """Build per-file records in one pass. If `roots` and `tracked_per_root`
    are supplied, each record is tagged with its containing git repo (if any)
    and whether the file is tracked in that repo.

    Cross-file enrichment (used_by, resolved dependencies) is NOT applied
    here — call enrich_cross_file_metadata once on the combined record list
    so tracked/untracked links are visible."""
    result: List[Tuple[Path, Dict[str, object]]] = []
    roots_list = list(roots) if roots else []
    tracked_map = tracked_per_root or {}
    for path in files:
        record = get_file_record(
            path=path,
            rel_path=path.relative_to(repo).as_posix(),
            preferred_encoding=preferred_encoding,
            max_text_file_kb=max_text_file_kb,
            include_no_extension=include_no_extension,
        )
        # Tag with git origin (which repo the file lives in, if any) and
        # whether that repo lists the file as tracked.
        if roots_list:
            origin_root, rel_in_root = assign_git_origin(path, roots_list)
            record["git_origin"] = origin_root
            if origin_root is not None:
                record["git_tracked"] = rel_in_root in tracked_map.get(origin_root, set())
        result.append((path, record))
    return result


def iter_all_records(
    tracked: Sequence[Tuple[Path, Dict[str, object]]],
    untracked: Sequence[Tuple[Path, Dict[str, object]]],
    local_only: Sequence[Tuple[Path, Dict[str, object]]],
) -> List[Tuple[Path, Dict[str, object]]]:
    """Flat union of the three record buckets. Each exporter computes the
    same union for header stats; centralizing here keeps the order stable
    (tracked → untracked → local-only) and the call site one line."""
    return list(tracked) + list(untracked) + list(local_only)


def total_content_tokens(
    records: Sequence[Tuple[Path, Dict[str, object]]],
) -> int:
    """Sum of per-record `tokens_estimate`. Returns 0 when annotation hasn't
    been run (every record has 0 by default)."""
    return sum(int(r.get("tokens_estimate", 0)) for _, r in records)


def build_file_heading(rel: str, record: Dict[str, object]) -> str:
    """One-line file heading. Always includes size + mtime; lines and token
    estimate are added when known. Used by docx, odt, and the md heading
    helper so every format sees the same shape."""
    bits: List[str] = [str(record["size_human"])]
    line_count = record.get("line_count")
    if line_count is not None:
        bits.append(f"{line_count} lines")
    tokens = int(record.get("tokens_estimate", 0))
    if tokens > 0:
        bits.append(f"~{tokens:,} tokens")
    bits.append(str(record["mtime_utc"]))
    return f"{rel} ({', '.join(bits)})"


def group_records_by_directory(
    repo: Path,
    records: Sequence[Tuple[Path, Dict[str, object]]],
) -> List[Tuple[str, List[Tuple[Path, Dict[str, object]]]]]:
    grouped: Dict[str, List[Tuple[Path, Dict[str, object]]]] = defaultdict(list)
    for path, record in records:
        rel = path.relative_to(repo).as_posix()
        bucket = infer_directory_bucket(rel)
        grouped[bucket].append((path, record))
    return sorted(grouped.items(), key=lambda x: (x[0] != "(repo root)", x[0].lower()))


def resolve_include_candidate_paths(rel_parent: Path, include_path: str) -> List[Path]:
    inc = Path(include_path.replace("\\", "/"))
    candidates: List[str] = []
    if not inc.is_absolute():
        candidates.append((rel_parent / inc).as_posix())
        candidates.append(inc.as_posix())
        if not inc.as_posix().startswith("src/"):
            candidates.append(("src/" + inc.as_posix()).replace("//", "/"))

    deduped: List[Path] = []
    seen = set()
    for candidate in candidates:
        normalized = candidate.replace("\\", "/")
        if normalized not in seen:
            seen.add(normalized)
            deduped.append(Path(normalized))
    return deduped


def enrich_cross_file_metadata(repo: Path, records: Sequence[Tuple[Path, Dict[str, object]]]) -> None:
    rel_to_record: Dict[str, Dict[str, object]] = {}
    basename_index: Dict[str, List[str]] = defaultdict(list)

    for path, record in records:
        rel = path.relative_to(repo).as_posix()
        rel_to_record[rel] = record
        basename_index[path.name].append(rel)

    reverse_edges: Dict[str, Set[str]] = defaultdict(set)

    for path, record in records:
        rel = path.relative_to(repo).as_posix()
        rel_parent = Path(rel).parent
        resolved: List[str] = []
        for inc in record.get("dependencies_raw", []):
            found = None
            for candidate in resolve_include_candidate_paths(rel_parent, inc):
                c = candidate.as_posix().replace("\\", "/")
                if c in rel_to_record:
                    found = c
                    break
            if found is None:
                base = Path(inc).name
                base_matches = basename_index.get(base, [])
                if len(base_matches) == 1:
                    found = base_matches[0]

            if found:
                label = f"{include_path_to_module_name(inc)} ({found})"
                reverse_edges[found].add(rel)

                stem = Path(found).stem
                suffix = Path(found).suffix.lower()
                if suffix in {".hpp", ".hh", ".h"}:
                    parent = str(Path(found).parent).replace("\\", "/")
                    for ext in [".cpp", ".cc", ".cxx", ".c"]:
                        impl = f"{parent}/{stem}{ext}".replace("//", "/")
                        if impl in rel_to_record:
                            if impl != rel:
                                reverse_edges[impl].add(rel)
            else:
                label = include_path_to_module_name(inc)

            if label not in resolved:
                resolved.append(label)

        record["dependencies_resolved"] = resolved[:12]

    for path, record in records:
        rel = path.relative_to(repo).as_posix()
        users = sorted(reverse_edges.get(rel, set()))
        record["used_by"] = [Path(user).name for user in users[:12]]


def describe_directory_bucket(bucket_records: Sequence[Tuple[Path, Dict[str, object]]]) -> str:
    """Auto-derive a one-line label: file count, total lines, dominant language."""
    file_count = len(bucket_records)
    line_count = sum(int(r.get("line_count") or 0) for _, r in bucket_records)
    type_counts: Dict[str, int] = defaultdict(int)
    for _, record in bucket_records:
        type_counts[str(record.get("type", "Unknown"))] += 1
    top_type = max(type_counts.items(), key=lambda kv: kv[1])[0] if type_counts else "mixed"
    if line_count > 0:
        return f"{file_count} files, {line_count:,} lines, mostly {top_type}"
    return f"{file_count} files, mostly {top_type}"


def add_architecture_summary(
    doc,
    repo: Path,
    has_git: bool,
    tracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    untracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    local_only_records: Sequence[Tuple[Path, Dict[str, object]]],
) -> None:
    # Use the full union — in workspace mode (has_git=True but with no-repo
    # files alongside subrepos), the old "tracked + untracked or local-only"
    # split would silently drop the no-repo files from the directory groups.
    records = iter_all_records(tracked_records, untracked_records, local_only_records)
    grouped = group_records_by_directory(repo, records)

    p = doc.add_heading("Architecture", level=1)
    compact_paragraph(p, after_pt=2, before_pt=3)

    add_compact_paragraph(doc, "Directory groups:", after_pt=1)
    for bucket, bucket_records in grouped:
        add_compact_paragraph(
            doc,
            f"- {bucket}: {describe_directory_bucket(bucket_records)}",
            after_pt=0,
        )


def add_recent_files_block(doc, title: str, records: Sequence[Tuple[Path, Dict[str, object]]], repo: Path, top_n: int = 5):
    p = doc.add_heading(title, level=1)
    compact_paragraph(p, after_pt=2, before_pt=3)

    if not records:
        add_compact_paragraph(doc, "(none)", after_pt=2)
        return

    add_compact_paragraph(doc, "Sorted by modification time, newest first:", after_pt=1)
    top_records = sorted(records, key=lambda item: str(item[1]["mtime_utc"]), reverse=True)[:top_n]
    for path, record in top_records:
        rel = path.relative_to(repo).as_posix()
        add_compact_paragraph(doc, f"- {rel} ({record['mtime_utc']})", after_pt=0)


def collect_ranked_entry_points(records: Sequence[Tuple[Path, Dict[str, object]]], repo: Path) -> List[Tuple[int, str, str, str]]:
    ranked: List[Tuple[int, str, str, str]] = []
    confidence_score = {"high": 3, "medium": 2, "low": 1}
    for path, record in records:
        rel = path.relative_to(repo).as_posix()
        for label, confidence in record.get("entry_signals", []):
            ranked.append((confidence_score.get(confidence, 0), rel, label, confidence))
    return sorted(ranked, key=lambda x: (-x[0], x[1], x[2]))


def harvest_project_symbols(
    records: Sequence[Tuple[Path, Dict[str, object]]],
    repo: Path,
) -> Tuple[Dict[str, str], Dict[str, List[str]]]:
    """Map every project-defined function name to its file. Used to filter
    out external/library calls so the trace only follows project code."""
    sym_to_file: Dict[str, str] = {}
    short_to_qual: Dict[str, List[str]] = defaultdict(list)
    for path, record in records:
        if not record.get("content_included"):
            continue
        suffix = path.suffix.lower()
        patterns = FUNCTION_DEF_PATTERNS_BY_SUFFIX.get(suffix)
        if not patterns:
            continue
        rel = path.relative_to(repo).as_posix()
        text = str(record.get("content_or_note", ""))
        for pattern in patterns:
            for match in pattern.finditer(text):
                qual = match.group(1)
                if qual in CONTROL_FLOW_NAMES:
                    continue
                sym_to_file.setdefault(qual, rel)
                if "::" in qual:
                    short = qual.rsplit("::", 1)[-1]
                    if rel not in short_to_qual[short]:
                        short_to_qual[short].append(rel)
                    sym_to_file.setdefault(short, rel)
    return sym_to_file, short_to_qual


def _strip_for_call_detection(text: str, suffix: str) -> str:
    """Internal helper for the call-graph sketch: strips BOTH string literals
    and comments so a regex hunting for `name(` calls inside a function body
    does not match identifiers that appear inside strings or comments.

    Destructive — do NOT use for user-facing content. The user-facing
    counterpart is `strip_comments_only`, which preserves strings."""
    out: List[str] = []
    i = 0
    n = len(text)
    py_like = suffix in {".py", ".pyw"}
    sh_like = suffix in {".sh", ".rb"}
    cstyle = suffix not in {".py", ".pyw"}
    while i < n:
        c = text[i]
        nxt = text[i + 1] if i + 1 < n else ""
        if cstyle and c == "/" and nxt == "/":
            j = text.find("\n", i)
            if j < 0:
                break
            out.append(" " * (j - i))
            i = j
            continue
        if cstyle and c == "/" and nxt == "*":
            j = text.find("*/", i + 2)
            if j < 0:
                break
            out.append(" " * (j + 2 - i))
            i = j + 2
            continue
        if (py_like or sh_like) and c == "#":
            j = text.find("\n", i)
            if j < 0:
                break
            out.append(" " * (j - i))
            i = j
            continue
        if py_like and (text.startswith('"""', i) or text.startswith("'''", i)):
            quote = text[i:i + 3]
            j = text.find(quote, i + 3)
            if j < 0:
                break
            out.append(" " * (j + 3 - i))
            i = j + 3
            continue
        if c in ('"', "'"):
            quote = c
            j = i + 1
            while j < n:
                if text[j] == "\\":
                    j += 2
                    continue
                if text[j] == quote:
                    break
                if text[j] == "\n" and not py_like:
                    break
                j += 1
            out.append(" " * (j + 1 - i))
            i = j + 1
            continue
        out.append(c)
        i += 1
    return "".join(out)


def extract_function_body(text: str, name: str, suffix: str) -> str:
    """Best-effort body extraction. Python uses indent rules; brace-based
    languages walk balanced braces. Returns '' if the function cannot be located."""
    suffix = suffix.lower()
    patterns = FUNCTION_DEF_PATTERNS_BY_SUFFIX.get(suffix)
    if not patterns:
        return ""

    short_name = name.rsplit("::", 1)[-1]
    candidate: Optional[Tuple[int, int]] = None
    for pattern in patterns:
        for match in pattern.finditer(text):
            captured = match.group(1)
            if captured == name or captured == short_name:
                candidate = (match.start(), match.end())
                break
        if candidate:
            break
    if not candidate:
        return ""

    if suffix in {".py", ".pyw"}:
        line_start = text.rfind("\n", 0, candidate[0]) + 1
        signature_indent = 0
        for ch in text[line_start:candidate[0] + 1]:
            if ch == " ":
                signature_indent += 1
            elif ch == "\t":
                signature_indent += 4
            else:
                break
        body_start = text.find(":", candidate[1])
        if body_start < 0:
            return ""
        body_start = text.find("\n", body_start)
        if body_start < 0:
            return ""
        body_start += 1
        lines: List[str] = []
        for raw_line in text[body_start:].splitlines(keepends=True):
            stripped = raw_line.strip()
            if not stripped:
                lines.append(raw_line)
                continue
            indent = 0
            for ch in raw_line:
                if ch == " ":
                    indent += 1
                elif ch == "\t":
                    indent += 4
                else:
                    break
            if indent <= signature_indent:
                break
            lines.append(raw_line)
        return "".join(lines)

    open_brace = text.find("{", candidate[1] - 1)
    if open_brace < 0:
        return ""
    depth = 0
    i = open_brace
    n = len(text)
    py_like = False
    while i < n:
        ch = text[i]
        nxt = text[i + 1] if i + 1 < n else ""
        if ch == "/" and nxt == "/":
            j = text.find("\n", i)
            if j < 0:
                break
            i = j
            continue
        if ch == "/" and nxt == "*":
            j = text.find("*/", i + 2)
            if j < 0:
                break
            i = j + 2
            continue
        if ch in ('"', "'"):
            quote = ch
            j = i + 1
            while j < n:
                if text[j] == "\\":
                    j += 2
                    continue
                if text[j] == quote:
                    break
                if text[j] == "\n":
                    break
                j += 1
            i = j + 1
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[open_brace + 1:i]
        i += 1
    return ""


def find_calls_in_body(body: str, suffix: str, sym_to_file: Dict[str, str]) -> List[str]:
    """Return ordered, de-duplicated callee names that resolve to project symbols."""
    cleaned = _strip_for_call_detection(body, suffix)
    seen: Set[str] = set()
    ordered: List[str] = []
    for match in CALL_REGEX.finditer(cleaned):
        name = match.group(1)
        if name in CONTROL_FLOW_NAMES:
            continue
        target_file = sym_to_file.get(name)
        if not target_file:
            short = name.rsplit("::", 1)[-1] if "::" in name else None
            if short:
                target_file = sym_to_file.get(short)
        if not target_file:
            continue
        if name in seen:
            continue
        seen.add(name)
        ordered.append(name)
    return ordered


def build_entry_point_trace(
    records: Sequence[Tuple[Path, Dict[str, object]]],
    repo: Path,
    max_entries: int = 5,
    max_first_level: int = 6,
    max_second_level: int = 4,
) -> List[Dict[str, object]]:
    """For the top-confidence entry points, trace one to two levels of
    project-internal calls. External/library calls are excluded by filtering
    against project-defined symbols."""
    sym_to_file, _ = harvest_project_symbols(records, repo)
    if not sym_to_file:
        return []

    text_by_rel: Dict[str, str] = {}
    suffix_by_rel: Dict[str, str] = {}
    for path, record in records:
        if not record.get("content_included"):
            continue
        rel = path.relative_to(repo).as_posix()
        text_by_rel[rel] = str(record.get("content_or_note", ""))
        suffix_by_rel[rel] = path.suffix.lower()

    ranked = collect_ranked_entry_points(records, repo)
    picked: List[Tuple[str, str, str]] = []
    seen_rel: Set[str] = set()
    for _, rel, label, confidence in ranked:
        if rel in seen_rel:
            continue
        seen_rel.add(rel)
        picked.append((rel, label, confidence))
        if len(picked) >= max_entries:
            break

    traces: List[Dict[str, object]] = []
    for rel, label, confidence in picked:
        text = text_by_rel.get(rel, "")
        suffix = suffix_by_rel.get(rel, "")
        if not text or not suffix:
            continue
        seed_name = _entry_seed_name(label, text, suffix)
        if not seed_name:
            continue
        body = extract_function_body(text, seed_name, suffix)
        if not body:
            continue
        first_calls = find_calls_in_body(body, suffix, sym_to_file)[:max_first_level]
        children: List[Dict[str, object]] = []
        for callee in first_calls:
            target_rel = sym_to_file.get(callee) or sym_to_file.get(callee.rsplit("::", 1)[-1])
            grandchildren: List[str] = []
            if target_rel and target_rel in text_by_rel:
                sub_body = extract_function_body(
                    text_by_rel[target_rel], callee, suffix_by_rel[target_rel]
                )
                if sub_body:
                    grandchildren = find_calls_in_body(
                        sub_body, suffix_by_rel[target_rel], sym_to_file
                    )[:max_second_level]
            children.append({
                "name": callee,
                "file": target_rel or "",
                "calls": grandchildren,
            })
        traces.append({
            "rel": rel,
            "label": label,
            "confidence": confidence,
            "seed": seed_name,
            "calls": children,
        })
    return traces


def _entry_seed_name(label: str, text: str, suffix: str) -> str:
    """Derive a function name to start tracing from. Prefer 'main' when the
    entry-point label suggests it; otherwise pick the first defined function."""
    label_lower = label.lower()
    candidates: List[str] = []
    if "main" in label_lower:
        candidates.append("main")
    if "wmain" in label_lower:
        candidates.append("wmain")
    if "winmain" in label_lower or "winmain" in text.lower():
        candidates.append("WinMain")
    patterns = FUNCTION_DEF_PATTERNS_BY_SUFFIX.get(suffix, [])
    defined: List[str] = []
    seen_names: Set[str] = set()
    for pattern in patterns:
        for match in pattern.finditer(text):
            name = match.group(1)
            if name in CONTROL_FLOW_NAMES or name in seen_names:
                continue
            seen_names.add(name)
            defined.append(name)
            if len(defined) >= 128:
                break
        if len(defined) >= 128:
            break
    for candidate in candidates:
        if candidate in defined:
            return candidate
        for d in defined:
            if d.rsplit("::", 1)[-1] == candidate:
                return d
    return defined[0] if defined else ""


def add_entry_trace_block(
    doc,
    title: str,
    records: Sequence[Tuple[Path, Dict[str, object]]],
    repo: Path,
) -> None:
    p = doc.add_heading(title, level=1)
    compact_paragraph(p, after_pt=2, before_pt=3)

    traces = build_entry_point_trace(records, repo)
    if not traces:
        add_compact_paragraph(doc, "(no traceable entry points)", after_pt=2)
        return

    add_compact_paragraph(
        doc,
        "Static call sketch — project-defined callees only, library calls omitted. "
        "Two levels deep, regex-based, best-effort.",
        after_pt=1,
    )
    for trace in traces:
        add_compact_paragraph(
            doc,
            f"- {trace['rel']} :: {trace['seed']}() [{trace['confidence']}] — {trace['label']}",
            after_pt=0,
        )
        children = trace.get("calls", []) or []
        if not children:
            add_compact_paragraph(doc, "    (no project-internal calls detected)", after_pt=0)
            continue
        for child in children:
            child_file = f" → {child['file']}" if child.get("file") else ""
            add_compact_paragraph(doc, f"    → {child['name']}(){child_file}", after_pt=0)
            for grand in child.get("calls", []) or []:
                add_compact_paragraph(doc, f"        → {grand}()", after_pt=0)


def add_entry_points_block(doc, title: str, records: Sequence[Tuple[Path, Dict[str, object]]], repo: Path, top_n: int = 12):
    p = doc.add_heading(title, level=1)
    compact_paragraph(p, after_pt=2, before_pt=3)

    ranked = collect_ranked_entry_points(records, repo)
    if not ranked:
        add_compact_paragraph(doc, "(no strong entry-point signals detected)", after_pt=2)
        return

    app_like = []
    test_like = []
    for item in ranked:
        rel = item[1]
        if rel.startswith("tests/") or "/test" in rel or rel.endswith("_test.cpp") or rel.endswith("_tests.cpp"):
            test_like.append(item)
        else:
            app_like.append(item)

    def emit_group(group_title: str, items: List[Tuple[int, str, str, str]]):
        add_compact_paragraph(doc, group_title + ":", after_pt=1)
        seen = set()
        count = 0
        for _, rel, label, confidence in items:
            key = (rel, label)
            if key in seen:
                continue
            seen.add(key)
            add_compact_paragraph(doc, f"- {rel} [{confidence}] — {label}", after_pt=0)
            count += 1
            if count >= top_n:
                break
        if count == 0:
            add_compact_paragraph(doc, "(none)", after_pt=0)

    emit_group("Application entry points", app_like)
    emit_group("Test entry points", test_like)


def add_core_files_block(doc, title: str, records: Sequence[Tuple[Path, Dict[str, object]]], repo: Path, top_n: int = 8):
    p = doc.add_heading(title, level=1)
    compact_paragraph(p, after_pt=2, before_pt=3)

    if not records:
        add_compact_paragraph(doc, "(none)", after_pt=2)
        return

    ranked = []
    for path, record in records:
        rel = path.relative_to(repo).as_posix()
        role = str(record.get("role_hint", ""))
        line_count = int(record.get("line_count") or 0)
        used_by = len(record.get("used_by", []))
        deps = len(record.get("dependencies_resolved", []))
        bonus = 0
        if role == "code":
            bonus += 20
        if role in {"test", "docs", "build", "asset", "config"}:
            bonus -= 50
        score = (line_count // 20) + (used_by * 8) + (deps * 2) + bonus
        ranked.append((score, rel, role, line_count, used_by, deps))

    for _, rel, role, line_count, used_by, deps in sorted(ranked, key=lambda x: (-x[0], x[1]))[:top_n]:
        add_compact_paragraph(
            doc,
            f"- {rel} [{role}] — {line_count} lines, used_by: {used_by}, deps: {deps}",
            after_pt=0,
        )


# ---------------- DOCX formatting helpers ----------------

def compact_paragraph(paragraph, after_pt: float = 0, before_pt: float = 0, line_spacing: float = 1.0) -> None:
    from docx.shared import Pt
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing = line_spacing


def add_compact_paragraph(doc, text: str = "", style: Optional[str] = None, after_pt: float = 0, before_pt: float = 0):
    p = doc.add_paragraph(text, style=style)
    compact_paragraph(p, after_pt=after_pt, before_pt=before_pt, line_spacing=1.0)
    return p


def add_docx_plain_marker_paragraph(doc, text: str, font_size_pt: int = 8, after_pt: float = 0):
    from docx.shared import Pt, RGBColor

    para = doc.add_paragraph()
    compact_paragraph(para, after_pt=after_pt, before_pt=0, line_spacing=1.0)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = RGBColor(110, 110, 110)
    return para


def extract_readme_excerpt(
    records: Sequence[Tuple[Path, Dict[str, object]]],
    repo: Path,
    max_chars: int = 500,
) -> Optional[str]:
    candidates = ("readme.md", "readme.rst", "readme.txt", "readme")
    for path, record in records:
        rel = path.relative_to(repo).as_posix().lower()
        if rel in candidates and bool(record.get("content_included")):
            content = str(record.get("content_or_note", "")).strip()
            if not content:
                continue
            if len(content) <= max_chars:
                return content
            cut = content[:max_chars]
            last_break = max(cut.rfind("\n\n"), cut.rfind(". "))
            if last_break > max_chars // 2:
                cut = cut[:last_break].rstrip()
            return cut + " …"
    return None


def add_at_a_glance_block(
    doc,
    summary: Dict[str, object],
    all_records: Sequence[Tuple[Path, Dict[str, object]]],
    repo: Path,
    has_git: bool,
    tracked_summary: Optional[Dict[str, object]] = None,
    untracked_summary: Optional[Dict[str, object]] = None,
) -> None:
    p = doc.add_heading("At a glance", level=1)
    compact_paragraph(p, after_pt=2, before_pt=3)

    file_count = int(summary["file_count"])
    text_count = int(summary["text_count"])
    binary_count = int(summary["binary_count"])
    total_size = int(summary["total_size_bytes"])
    lines_total = int(summary["lines_total"])
    lines_by_kind = summary["lines_by_kind"]
    lines_by_language = summary["lines_by_language"]

    add_compact_paragraph(
        doc,
        f"Files: {file_count} ({text_count} text, {binary_count} binary). "
        f"Total size: {format_file_size_kb(total_size)}.",
        after_pt=0,
    )

    if has_git and tracked_summary is not None and untracked_summary is not None:
        t_files = int(tracked_summary["file_count"])
        t_size = format_file_size_kb(int(tracked_summary["total_size_bytes"]))
        u_files = int(untracked_summary["file_count"])
        u_size = format_file_size_kb(int(untracked_summary["total_size_bytes"]))
        add_compact_paragraph(
            doc,
            f"Git split: tracked {t_files} files / {t_size}; untracked {u_files} files / {u_size}.",
            after_pt=1,
        )
    else:
        add_compact_paragraph(doc, "", after_pt=1)

    add_compact_paragraph(doc, f"Total lines (text files): {lines_total:,}", after_pt=0)
    for kind in ("code", "docs", "data", "config", "other"):
        n = int(lines_by_kind.get(kind, 0))
        if n > 0:
            add_compact_paragraph(doc, f"  {kind}: {n:,} lines", after_pt=0)
    add_compact_paragraph(doc, "", after_pt=1)

    top_langs = sorted(lines_by_language.items(), key=lambda kv: -int(kv[1]))[:6]
    if top_langs:
        add_compact_paragraph(doc, "Top languages by lines:", after_pt=0)
        for lang, n in top_langs:
            if int(n) > 0:
                add_compact_paragraph(doc, f"  {lang}: {int(n):,}", after_pt=0)
        add_compact_paragraph(doc, "", after_pt=1)

    if all_records:
        sorted_by_lines = sorted(
            all_records,
            key=lambda item: int(item[1].get("line_count") or 0),
            reverse=True,
        )
        path, record = sorted_by_lines[0]
        rel = path.relative_to(repo).as_posix()
        n = record.get("line_count")
        if isinstance(n, int) and n > 0:
            add_compact_paragraph(
                doc,
                f"Largest file: {rel} ({n:,} lines, {record['size_human']})",
                after_pt=1,
            )

    excerpt = extract_readme_excerpt(all_records, repo)
    if excerpt:
        add_compact_paragraph(doc, "README excerpt:", after_pt=0)
        for line in excerpt.splitlines() or [""]:
            add_compact_paragraph(doc, line, after_pt=0)
        add_compact_paragraph(doc, "", after_pt=1)


def add_unified_toc(
    doc,
    title: str,
    tracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    untracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    local_only_records: Sequence[Tuple[Path, Dict[str, object]]],
    repo: Path,
    has_git: bool,
    git_roots: Optional[Sequence[Tuple[Path, str]]] = None,
) -> None:
    p = doc.add_heading(title, level=1)
    compact_paragraph(p, after_pt=2, before_pt=3)

    # Render every record in canonical order. The marker (T/U/N or
    # T:subrepo/U:subrepo) is derived per-record from git_origin so
    # workspaces with multiple subrepos show which repo each file belongs to.
    all_recs = (list(tracked_records) + list(untracked_records)
                + list(local_only_records))
    if not all_recs:
        add_compact_paragraph(doc, "(none)", after_pt=2)
        return

    multi_root = bool(git_roots and len(git_roots) > 1)
    legend_bits: List[str] = []
    if any(r.get("git_origin") is not None and r.get("git_tracked")
           for _, r in all_recs):
        legend_bits.append("T = tracked")
    if any(r.get("git_origin") is not None and not r.get("git_tracked")
           for _, r in all_recs):
        legend_bits.append("U = untracked")
    if any(r.get("git_origin") is None for _, r in all_recs):
        legend_bits.append("N = no-repo")
    if legend_bits:
        add_compact_paragraph(doc, ", ".join(legend_bits) + ".", after_pt=1)

    grouped: Dict[str, List[Tuple[Path, Dict[str, object]]]] = defaultdict(list)
    for path, record in all_recs:
        rel = path.relative_to(repo).as_posix()
        grouped[infer_directory_bucket(rel)].append((path, record))

    for bucket in sorted(grouped.keys(), key=lambda b: (b != "(repo root)", b.lower())):
        p_bucket = doc.add_heading(bucket, level=2)
        compact_paragraph(p_bucket, after_pt=1, before_pt=2)
        for path, record in grouped[bucket]:
            rel = path.relative_to(repo).as_posix()
            size_human = str(record["size_human"])
            line_count = record["line_count"]
            role = str(record.get("role_hint", ""))
            type_label = str(record.get("type", ""))
            lines_str = f"{line_count} lines" if line_count is not None else "—"
            tokens = int(record.get("tokens_estimate", 0))
            tokens_str = f"~{tokens:,} tok" if tokens > 0 else "— tok"
            tag = render_toc_marker(record, repo, multi_root)
            label = (f"{tag}  {rel}  |  {size_human}  |  {lines_str}  |  "
                     f"{tokens_str}  |  {type_label}  |  role:{role}")
            p = doc.add_paragraph(label, style="List Bullet")
            compact_paragraph(p, after_pt=0, before_pt=0)


def add_structured_file_header(doc, rel: str, record: Dict[str, object]) -> None:
    complexity = record.get("complexity", {})
    add_docx_plain_marker_paragraph(doc, "--- FILE ---", after_pt=0)
    add_docx_plain_marker_paragraph(doc, f"path: {rel}", after_pt=0)
    add_docx_plain_marker_paragraph(doc, f"type: {record['type']}", after_pt=0)
    add_docx_plain_marker_paragraph(doc, f"role: {record['role_hint']}", after_pt=0)
    add_docx_plain_marker_paragraph(doc, f"size: {record['size_human']}", after_pt=0)
    add_docx_plain_marker_paragraph(doc, f"mtime: {record['mtime_utc']}", after_pt=0)
    add_docx_plain_marker_paragraph(doc, f"id: sha1:{record['quick_hash']}", after_pt=0)
    if record.get("line_count") is not None:
        add_docx_plain_marker_paragraph(doc, f"lines: {record['line_count']}", after_pt=0)
    tokens = int(record.get("tokens_estimate", 0))
    if tokens > 0:
        add_docx_plain_marker_paragraph(doc, f"tokens: ~{tokens:,}", after_pt=0)
    classes = int(complexity.get("classes", 0))
    structs = int(complexity.get("structs", 0))
    functions_known = bool(complexity.get("functions_known", False))
    functions = int(complexity.get("functions", 0)) if functions_known else 0
    if classes or structs or functions:
        bits = []
        if functions:
            bits.append(f"functions={functions}")
        if classes:
            bits.append(f"classes={classes}")
        if structs:
            bits.append(f"structs={structs}")
        add_docx_plain_marker_paragraph(doc, "decls: " + ", ".join(bits), after_pt=0)
    deps = record.get("dependencies_resolved", [])
    add_docx_plain_marker_paragraph(doc, f"deps: {', '.join(deps) if deps else '(none)'}", after_pt=0)
    used_by = record.get("used_by", [])
    add_docx_plain_marker_paragraph(doc, f"used_by: {', '.join(used_by) if used_by else '(none)'}", after_pt=0)
    signals = [f"{label} [{confidence}]" for label, confidence in record.get("entry_signals", [])]
    add_docx_plain_marker_paragraph(doc, f"entry_signals: {', '.join(signals) if signals else '(none)'}", after_pt=1)


def add_docx_file_entries(
    doc,
    section_title: str,
    records: Sequence[Tuple[Path, Dict[str, object]]],
    repo: Path,
) -> None:
    from docx.shared import Pt

    doc.add_page_break()
    p = doc.add_heading(section_title, level=1)
    compact_paragraph(p, after_pt=3, before_pt=0)

    if not records:
        add_compact_paragraph(doc, "(none)", after_pt=2)
        return

    grouped = group_records_by_directory(repo, records)

    for bucket, bucket_records in grouped:
        p_bucket = doc.add_heading(bucket, level=2)
        compact_paragraph(p_bucket, after_pt=2, before_pt=3)

        for path, record in bucket_records:
            rel = path.relative_to(repo).as_posix()
            heading_text = build_file_heading(rel, record)
            p = doc.add_heading(heading_text, level=3)
            compact_paragraph(p, after_pt=1, before_pt=2)

            add_structured_file_header(doc, rel, record)

            if record["content_included"]:
                add_docx_plain_marker_paragraph(doc, "--- BEGIN CONTENT ---", after_pt=1)

                content_para = doc.add_paragraph()
                compact_paragraph(content_para, after_pt=1, before_pt=0)
                run = content_para.add_run(sanitize_xml_compatible_text(str(record["content_or_note"])))
                run.font.name = "Consolas"
                run.font.size = Pt(8)

                add_docx_plain_marker_paragraph(doc, "--- END FILE ---", after_pt=2)
            else:
                add_docx_plain_marker_paragraph(
                    doc,
                    f"omission: {record['omission_reason'] or 'binary or non-text file'}",
                    after_pt=0,
                )
                add_docx_plain_marker_paragraph(doc, f"note: {record['content_or_note']}", after_pt=0)
                add_docx_plain_marker_paragraph(doc, "--- END FILE ---", after_pt=2)


def export_docx(
    repo: Path,
    tracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    untracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    local_only_records: Sequence[Tuple[Path, Dict[str, object]]],
    output_path: Path,
    has_git: bool,
    branch_name: str,
    head_commit: str,
    head_commit_datetime: str,
    report_created: str,
    sections: Optional[Set[str]] = None,
    token_model: str = "gpt-4o",
    git_roots: Optional[Sequence[Tuple[Path, str]]] = None,
) -> None:
    if sections is None:
        sections = set(DEFAULT_SECTIONS)
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency for Word export. Install with: py -m pip install python-docx"
        ) from exc

    doc = Document()

    cp = doc.core_properties
    cp.author = generator_name()
    cp.last_modified_by = generator_name()

    p = doc.add_heading(f"Repository export: {repo.name}", level=0)
    compact_paragraph(p, after_pt=3, before_pt=0)

    content_tokens_total = total_content_tokens(
        iter_all_records(tracked_records, untracked_records, local_only_records)
    )

    add_compact_paragraph(doc, f"Root folder: {repo}", after_pt=0)
    add_compact_paragraph(doc, f"Report created (UTC): {report_created}", after_pt=0)
    add_compact_paragraph(
        doc,
        f"Content tokens (estimate, tokenizer {token_model}): ~{content_tokens_total:,}",
        after_pt=0,
    )
    add_compact_paragraph(
        doc,
        format_git_roots_summary(
            git_roots or [], repo,
            tracked_records, untracked_records, local_only_records,
        ),
        after_pt=0,
    )

    if has_git:
        add_compact_paragraph(doc, "Git repo detected: yes", after_pt=0)
        add_compact_paragraph(doc, f"Git branch: {branch_name}", after_pt=0)
        add_compact_paragraph(doc, f"HEAD commit: {head_commit}", after_pt=0)
        add_compact_paragraph(doc, f"HEAD commit date (UTC): {head_commit_datetime}", after_pt=2)
    else:
        add_compact_paragraph(doc, "Git repo detected: no", after_pt=2)

    # Unified record view — covers self-repo, workspace, ancestor, and
    # no-repo cases without three separate code paths. In workspace mode
    # tracked+untracked AND local-only records coexist; the old has_git
    # dichotomy would have silently dropped one side.
    all_records = iter_all_records(tracked_records, untracked_records, local_only_records)
    overall_summary = summarize_records([r for _, r in all_records])
    tracked_summary = (
        summarize_records([r for _, r in tracked_records]) if tracked_records else None
    )
    untracked_summary = (
        summarize_records([r for _, r in untracked_records]) if untracked_records else None
    )

    if "glance" in sections:
        add_at_a_glance_block(
            doc, overall_summary, all_records, repo, has_git,
            tracked_summary=tracked_summary, untracked_summary=untracked_summary,
        )
    if "recent" in sections:
        add_recent_files_block(doc, "Recent files", all_records, repo)
    if "architecture" in sections:
        add_architecture_summary(doc, repo, has_git, tracked_records, untracked_records, local_only_records)
    if "entry_points" in sections:
        add_entry_points_block(doc, "Entry points", all_records, repo)
    if "trace" in sections:
        add_entry_trace_block(doc, "Entry-point call trace", all_records, repo)
    if "core" in sections:
        add_core_files_block(doc, "Core files", all_records, repo)
    if "toc" in sections:
        add_unified_toc(
            doc, "Table of contents",
            tracked_records, untracked_records, local_only_records, repo, has_git,
            git_roots=git_roots,
        )
    if "entries" in sections:
        # Emit a section per non-empty bucket; in workspace mode all three
        # can be populated simultaneously.
        if tracked_records:
            add_docx_file_entries(doc, "Tracked file entries", tracked_records, repo)
        if untracked_records:
            add_docx_file_entries(doc, "Untracked file entries", untracked_records, repo)
        if local_only_records:
            title = "No-repo file entries" if has_git else "File entries"
            add_docx_file_entries(doc, title, local_only_records, repo)

    doc.save(output_path)


def export_xlsx(
    repo: Path,
    tracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    untracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    local_only_records: Sequence[Tuple[Path, Dict[str, object]]],
    output_path: Path,
    has_git: bool,
    branch_name: str,
    head_commit: str,
    head_commit_datetime: str,
    report_created: str,
    token_model: str = "gpt-4o",
    git_roots: Optional[Sequence[Tuple[Path, str]]] = None,
) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency for Excel export. Install with: py -m pip install openpyxl"
        ) from exc

    wb = Workbook()
    wb.properties.creator = generator_name()
    wb.properties.lastModifiedBy = generator_name()
    ws = wb.active
    ws.title = "Repository export"

    headers = [
        "section",          # tracked / untracked / local (= no-repo)
        "git_origin",       # subrepo label (workspace mode), empty otherwise
        "directory_group",
        "file_path",
        "type",
        "kind",
        "role",
        "size_bytes",
        "size_human",
        "mtime_utc",
        "quick_hash",
        "is_text",
        "line_count",
        "tokens_estimate",
        "functions_est",
        "classes_est",
        "structs_est",
        "dependencies",
        "used_by",
        "entry_signals",
        "content_included",
        "omission_reason",
        "content_or_note",
    ]
    ws.append(headers)

    header_font = Font(bold=True)
    wrap_alignment = Alignment(vertical="top", wrap_text=True)

    for col_idx in range(1, len(headers) + 1):
        ws.cell(row=1, column=col_idx).font = header_font

    # Iterate all three buckets unconditionally — in workspace mode (children
    # roots) both tracked/untracked and local-only records can be non-empty.
    groups = [
        ("tracked", tracked_records),
        ("untracked", untracked_records),
        ("local", local_only_records),
    ]

    for section_name, records in groups:
        for path, record in records:
            rel = path.relative_to(repo).as_posix()
            directory_group = infer_directory_bucket(rel)
            complexity = record.get("complexity", {})
            entry_signals = ", ".join(f"{label} [{confidence}]" for label, confidence in record.get("entry_signals", []))
            origin = record.get("git_origin")
            git_origin_label = render_repo_label(origin, repo) if origin else ""
            ws.append([
                section_name,
                git_origin_label,
                directory_group,
                rel,
                record["type"],
                record.get("kind", ""),
                record["role_hint"],
                record["size_bytes"],
                record["size_human"],
                record["mtime_utc"],
                record["quick_hash"],
                record["is_text"],
                record["line_count"] if record["line_count"] is not None else "",
                int(record.get("tokens_estimate", 0)) or "",
                complexity.get("functions", 0) if complexity.get("functions_known") else "",
                complexity.get("classes", 0),
                complexity.get("structs", 0),
                ", ".join(record.get("dependencies_resolved", [])),
                ", ".join(record.get("used_by", [])),
                entry_signals,
                record["content_included"],
                record["omission_reason"],
                sanitize_xml_compatible_text(str(record["content_or_note"])),
            ])

    ws.freeze_panes = "A2"
    # 23 columns: tokens_estimate at N (was M), git_origin newly inserted at B.
    widths = {
        "A": 14, "B": 16, "C": 20, "D": 60, "E": 20,
        "F": 14, "G": 14, "H": 14, "I": 16, "J": 18,
        "K": 10, "L": 12, "M": 12, "N": 14, "O": 12,
        "P": 12, "Q": 12, "R": 36, "S": 36, "T": 28,
        "U": 16, "V": 35, "W": 120,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    # Force every data row to a fixed height so 'content_or_note' (column U,
    # which can hold the full text of a source file) does not stretch the
    # row to thousands of points and turn the sheet into a wall of text.
    # Users can drag a row taller to inspect a single cell when needed.
    DATA_ROW_HEIGHT_PT = 30.0
    for row in ws.iter_rows(min_row=2):
        ws.row_dimensions[row[0].row].height = DATA_ROW_HEIGHT_PT
        for cell in row:
            cell.alignment = wrap_alignment

    content_tokens_total = total_content_tokens(
        iter_all_records(tracked_records, untracked_records, local_only_records)
    )

    meta = wb.create_sheet("Summary")
    meta["A1"] = "Repository"
    meta["B1"] = repo.name
    meta["A2"] = "Root folder"
    meta["B2"] = str(repo)
    meta["A3"] = "Report created (UTC)"
    meta["B3"] = report_created
    meta["A4"] = "Git repo detected"
    meta["B4"] = "yes" if has_git else "no"
    meta["A5"] = "Branch"
    meta["B5"] = branch_name if has_git else "n/a"
    meta["A6"] = "HEAD commit"
    meta["B6"] = head_commit if has_git else "n/a"
    meta["A7"] = "HEAD commit date (UTC)"
    meta["B7"] = head_commit_datetime if has_git else "n/a"
    meta["A8"] = "Content tokens (estimate)"
    meta["B8"] = content_tokens_total
    meta["A9"] = "Tokenizer model"
    meta["B9"] = token_model
    meta["A10"] = "Git roots summary"
    meta["B10"] = format_git_roots_summary(
        git_roots or [], repo,
        tracked_records, untracked_records, local_only_records,
    )
    # One row per discovered git root, listing kind + tracked/untracked counts.
    next_row = 11
    for root, kind in (git_roots or []):
        meta[f"A{next_row}"] = f"  · {render_repo_label(root, repo)}"
        t = sum(1 for _, r in tracked_records if r.get("git_origin") == root)
        u = sum(1 for _, r in untracked_records if r.get("git_origin") == root)
        meta[f"B{next_row}"] = f"{kind}; tracked={t}, untracked={u}"
        next_row += 1
    if local_only_records:
        meta[f"A{next_row}"] = "  · no-repo files"
        meta[f"B{next_row}"] = len(local_only_records)
    # A bit wider so the labels/values are readable.
    meta.column_dimensions["A"].width = 30
    meta.column_dimensions["B"].width = 80

    wb.save(output_path)


_MD_LANG_BY_TYPE: Dict[str, str] = {
    "Python": "python", "C": "c", "C++": "cpp", "C++ Header": "cpp",
    "C/C++ Header": "cpp", "C#": "csharp", "Java": "java",
    "JavaScript": "javascript", "JavaScript/JSX": "jsx",
    "TypeScript": "typescript", "TypeScript/TSX": "tsx",
    "Rust": "rust", "Go": "go", "Ruby": "ruby", "PHP": "php",
    "Kotlin": "kotlin", "Kotlin Script": "kotlin",
    "Scala": "scala", "Swift": "swift",
    "R": "r", "RMarkdown": "rmd",
    "HTML": "html", "CSS": "css", "SCSS": "scss", "Sass": "sass",
    "JSON": "json", "YAML": "yaml", "XML": "xml", "TOML": "toml",
    "Markdown": "markdown", "reStructuredText": "rst",
    "Shell": "bash", "PowerShell": "powershell",
    "SQL": "sql", "CMake": "cmake",
}


def _md_lang_hint(file_type: str) -> str:
    return _MD_LANG_BY_TYPE.get(file_type, "")


def _md_code_fence(content: str) -> str:
    """Return a backtick fence longer than any run of backticks in the content."""
    longest = 0
    cur = 0
    for ch in content:
        if ch == "`":
            cur += 1
            if cur > longest:
                longest = cur
        else:
            cur = 0
    return "`" * max(3, longest + 1)


def export_md(
    repo: Path,
    tracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    untracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    local_only_records: Sequence[Tuple[Path, Dict[str, object]]],
    output_path: Path,
    has_git: bool,
    branch_name: str,
    head_commit: str,
    head_commit_datetime: str,
    report_created: str,
    sections: Optional[Set[str]] = None,
    token_model: str = "gpt-4o",
    git_roots: Optional[Sequence[Tuple[Path, str]]] = None,
) -> None:
    if sections is None:
        sections = set(DEFAULT_SECTIONS)

    # Per-file token estimates (populated upstream by annotate_record_tokens).
    # We sum them here so the header carries a clear total even if the user
    # didn't pass --token-budget.
    content_tokens_total = total_content_tokens(
        iter_all_records(tracked_records, untracked_records, local_only_records)
    )
    have_tiktoken = _get_token_encoder(token_model) is not None
    token_label = "tokens" if have_tiktoken else "tokens (rough estimate)"

    out: List[str] = []
    out.append(f"<!-- generator: {generator_name()} -->")
    out.append(
        f"<!-- content {token_label}: ~{content_tokens_total:,} "
        f"(tokenizer: {token_model}) -->"
    )
    out.append(f"# Repository export: {repo.name}")
    out.append("")
    out.append(f"- Generator: {generator_name()}")
    out.append(f"- Root folder: `{repo}`")
    out.append(f"- Report created (UTC): {report_created}")
    out.append(
        f"- Content {token_label}: ~{content_tokens_total:,} "
        f"(tokenizer: {token_model})"
    )
    if has_git:
        out.append(f"- Git repo detected: yes")
        out.append(f"- Git branch: {branch_name}")
        out.append(f"- HEAD commit: {head_commit}")
        out.append(f"- HEAD commit date (UTC): {head_commit_datetime}")
    else:
        out.append("- Git repo detected: no")
    out.append(
        f"- {format_git_roots_summary(git_roots or [], repo, tracked_records, untracked_records, local_only_records)}"
    )
    out.append("")

    # Unified view; workspace mode (has_git=True with no-repo files alongside)
    # would otherwise drop the no-repo files. See export_docx for context.
    all_records = iter_all_records(tracked_records, untracked_records, local_only_records)
    overall_summary = summarize_records([r for _, r in all_records])
    tracked_summary = (
        summarize_records([r for _, r in tracked_records]) if tracked_records else None
    )
    untracked_summary = (
        summarize_records([r for _, r in untracked_records]) if untracked_records else None
    )

    if "glance" in sections:
        out.append("## At a glance")
        out.append("")
        out.append(
            f"- Files: {overall_summary['file_count']} "
            f"({overall_summary['text_count']} text, {overall_summary['binary_count']} binary)"
        )
        out.append(f"- Total size: {format_file_size_kb(int(overall_summary['total_size_bytes']))}")
        if has_git and tracked_summary and untracked_summary:
            out.append(
                f"- Git split: tracked {tracked_summary['file_count']} files / "
                f"{format_file_size_kb(int(tracked_summary['total_size_bytes']))}; "
                f"untracked {untracked_summary['file_count']} files / "
                f"{format_file_size_kb(int(untracked_summary['total_size_bytes']))}"
            )
        out.append(f"- Total lines (text files): {int(overall_summary['lines_total']):,}")
        for kind in ("code", "docs", "data", "config", "other"):
            n = int(overall_summary["lines_by_kind"].get(kind, 0))
            if n > 0:
                out.append(f"  - {kind}: {n:,} lines")
        top_langs = sorted(
            overall_summary["lines_by_language"].items(),
            key=lambda kv: -int(kv[1]),
        )[:6]
        if top_langs:
            out.append("- Top languages by lines:")
            for lang, n in top_langs:
                if int(n) > 0:
                    out.append(f"  - {lang}: {int(n):,}")
        if all_records:
            sorted_by_lines = sorted(
                all_records,
                key=lambda item: int(item[1].get("line_count") or 0),
                reverse=True,
            )
            path0, record0 = sorted_by_lines[0]
            n0 = record0.get("line_count")
            if isinstance(n0, int) and n0 > 0:
                out.append(
                    f"- Largest file: `{path0.relative_to(repo).as_posix()}` "
                    f"({n0:,} lines, {record0['size_human']})"
                )
        excerpt = extract_readme_excerpt(all_records, repo)
        if excerpt:
            out.append("")
            out.append("**README excerpt:**")
            out.append("")
            out.append("> " + excerpt.replace("\n", "\n> "))
        out.append("")

    if "recent" in sections and all_records:
        out.append("## Recent files")
        out.append("")
        out.append("Sorted by modification time, newest first:")
        out.append("")
        top_records = sorted(all_records, key=lambda item: str(item[1]["mtime_utc"]), reverse=True)[:5]
        for path, record in top_records:
            out.append(f"- `{path.relative_to(repo).as_posix()}` ({record['mtime_utc']})")
        out.append("")

    if "architecture" in sections:
        out.append("## Architecture")
        out.append("")
        if has_git:
            grouped = group_records_by_directory(repo, list(tracked_records) + list(untracked_records))
        else:
            grouped = group_records_by_directory(repo, local_only_records)
        for bucket, bucket_records in grouped:
            out.append(f"- **{bucket}** — {describe_directory_bucket(bucket_records)}")
        out.append("")

    if "entry_points" in sections:
        out.append("## Entry points")
        out.append("")
        ranked = collect_ranked_entry_points(all_records, repo)
        if not ranked:
            out.append("_no strong entry-point signals detected_")
        else:
            seen: Set[Tuple[str, str]] = set()
            for _, rel, label, confidence in ranked:
                key = (rel, label)
                if key in seen:
                    continue
                seen.add(key)
                out.append(f"- `{rel}` [{confidence}] — {label}")
        out.append("")

    if "trace" in sections:
        out.append("## Entry-point call trace")
        out.append("")
        traces = build_entry_point_trace(all_records, repo)
        if not traces:
            out.append("_no traceable entry points_")
        else:
            out.append(
                "_Static call sketch — project-defined callees only, library calls omitted. "
                "Two levels deep, regex-based, best-effort._"
            )
            out.append("")
            out.append("```")
            for trace in traces:
                out.append(f"{trace['rel']} :: {trace['seed']}()  [{trace['confidence']}]  {trace['label']}")
                children = trace.get("calls", []) or []
                if not children:
                    out.append("    (no project-internal calls detected)")
                    continue
                for child in children:
                    suffix = f"  → {child['file']}" if child.get("file") else ""
                    out.append(f"    → {child['name']}(){suffix}")
                    for grand in child.get("calls", []) or []:
                        out.append(f"        → {grand}()")
            out.append("```")
        out.append("")

    if "core" in sections:
        out.append("## Core files")
        out.append("")
        ranked: List[Tuple[int, str, str, int, int, int]] = []
        for path, record in all_records:
            rel = path.relative_to(repo).as_posix()
            role = str(record.get("role_hint", ""))
            line_count = int(record.get("line_count") or 0)
            used_by = len(record.get("used_by", []))
            deps = len(record.get("dependencies_resolved", []))
            bonus = 0
            if role == "code":
                bonus += 20
            if role in {"test", "docs", "build", "asset", "config"}:
                bonus -= 50
            score = (line_count // 20) + (used_by * 8) + (deps * 2) + bonus
            ranked.append((score, rel, role, line_count, used_by, deps))
        for _, rel, role, line_count, used_by, deps in sorted(ranked, key=lambda x: (-x[0], x[1]))[:8]:
            out.append(f"- `{rel}` [{role}] — {line_count} lines, used_by: {used_by}, deps: {deps}")
        out.append("")

    if "toc" in sections:
        out.append("## Table of contents")
        out.append("")
        multi_root = bool(git_roots and len(git_roots) > 1)

        def _toc_line(rel: str, record: Dict[str, object]) -> str:
            line_part = (f", {record['line_count']} lines"
                         if record.get("line_count") is not None else "")
            tok = int(record.get("tokens_estimate", 0))
            tok_part = f", ~{tok:,} tokens" if tok > 0 else ""
            marker = render_toc_marker(record, repo, multi_root)
            return (f"- [{marker}] `{rel}` ({record['size_human']}{line_part}"
                    f"{tok_part}, {record['mtime_utc']})")

        for path, record in iter_all_records(
            tracked_records, untracked_records, local_only_records
        ):
            out.append(_toc_line(path.relative_to(repo).as_posix(), record))
        out.append("")

    if "entries" in sections:
        # Emit a heading per non-empty bucket so workspace mode shows all
        # three (tracked + untracked subrepo files, plus no-repo siblings).
        if tracked_records:
            _md_emit_file_entries(out, "Tracked file entries", tracked_records, repo)
        if untracked_records:
            _md_emit_file_entries(out, "Untracked file entries", untracked_records, repo)
        if local_only_records:
            title = "No-repo file entries" if has_git else "File entries"
            _md_emit_file_entries(out, title, local_only_records, repo)

    output_path.write_text("\n".join(out).rstrip() + "\n", encoding="utf-8")


def _md_emit_file_entries(
    out: List[str],
    section_title: str,
    records: Sequence[Tuple[Path, Dict[str, object]]],
    repo: Path,
) -> None:
    out.append(f"## {section_title}")
    out.append("")
    if not records:
        out.append("_(none)_")
        out.append("")
        return
    grouped = group_records_by_directory(repo, records)
    for bucket, bucket_records in grouped:
        out.append(f"### {bucket}")
        out.append("")
        for path, record in bucket_records:
            rel = path.relative_to(repo).as_posix()
            out.append(f"#### `{rel}`")
            out.append("")
            out.append(f"- type: {record['type']}")
            out.append(f"- role: {record['role_hint']}")
            out.append(f"- size: {record['size_human']}")
            out.append(f"- mtime: {record['mtime_utc']}")
            out.append(f"- id: sha1:{record['quick_hash']}")
            if record.get("line_count") is not None:
                out.append(f"- lines: {record['line_count']}")
            tok = int(record.get("tokens_estimate", 0))
            if tok > 0:
                out.append(f"- tokens: ~{tok:,}")
            complexity = record.get("complexity", {})
            classes = int(complexity.get("classes", 0))
            structs = int(complexity.get("structs", 0))
            functions_known = bool(complexity.get("functions_known", False))
            functions = int(complexity.get("functions", 0)) if functions_known else 0
            if classes or structs or functions:
                bits: List[str] = []
                if functions:
                    bits.append(f"functions={functions}")
                if classes:
                    bits.append(f"classes={classes}")
                if structs:
                    bits.append(f"structs={structs}")
                out.append("- decls: " + ", ".join(bits))
            deps = record.get("dependencies_resolved", [])
            out.append(f"- deps: {', '.join(deps) if deps else '(none)'}")
            used_by = record.get("used_by", [])
            out.append(f"- used_by: {', '.join(used_by) if used_by else '(none)'}")
            signals = [f"{label} [{confidence}]" for label, confidence in record.get("entry_signals", [])]
            out.append(f"- entry_signals: {', '.join(signals) if signals else '(none)'}")
            out.append("")
            if record["content_included"]:
                content = sanitize_xml_compatible_text(str(record["content_or_note"]))
                fence = _md_code_fence(content)
                lang = _md_lang_hint(str(record["type"]))
                out.append(f"{fence}{lang}")
                out.append(content.rstrip("\n"))
                out.append(fence)
            else:
                out.append(f"_omission: {record['omission_reason'] or 'binary or non-text file'}_")
                out.append("")
                out.append(f"> {record['content_or_note']}")
            out.append("")


def export_json(
    repo: Path,
    tracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    untracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    local_only_records: Sequence[Tuple[Path, Dict[str, object]]],
    output_path: Path,
    has_git: bool,
    branch_name: str,
    head_commit: str,
    head_commit_datetime: str,
    report_created: str,
    token_model: str = "gpt-4o",
    git_roots: Optional[Sequence[Tuple[Path, str]]] = None,
) -> None:
    """Single JSON dump containing repo metadata, summary stats, and every
    file record. Always includes everything (no --sections gating)."""
    def _serialize_record(path: Path, record: Dict[str, object]) -> Dict[str, object]:
        rel = path.relative_to(repo).as_posix()
        signals = [
            {"label": label, "confidence": confidence}
            for label, confidence in record.get("entry_signals", [])
        ]
        complexity = dict(record.get("complexity", {}))
        origin = record.get("git_origin")
        return {
            "path": rel,
            "type": record["type"],
            "kind": record.get("kind", ""),
            "role_hint": record["role_hint"],
            "size_bytes": record["size_bytes"],
            "size_human": record["size_human"],
            "mtime_utc": record["mtime_utc"],
            "quick_hash": record["quick_hash"],
            "is_text": record["is_text"],
            "line_count": record["line_count"],
            "tokens_estimate": int(record.get("tokens_estimate", 0)),
            "git_origin": render_repo_label(origin, repo) if origin else None,
            "git_tracked": bool(record.get("git_tracked")),
            "complexity": complexity,
            "dependencies_raw": list(record.get("dependencies_raw", [])),
            "dependencies_resolved": list(record.get("dependencies_resolved", [])),
            "used_by": list(record.get("used_by", [])),
            "entry_signals": signals,
            "content_included": record["content_included"],
            "omission_reason": record["omission_reason"],
            "content_or_note": str(record["content_or_note"]),
        }

    def _serialize_summary(summary: Dict[str, object]) -> Dict[str, object]:
        return {
            "file_count": int(summary["file_count"]),
            "text_count": int(summary["text_count"]),
            "binary_count": int(summary["binary_count"]),
            "content_included_count": int(summary["content_included_count"]),
            "content_omitted_count": int(summary["content_omitted_count"]),
            "total_size_bytes": int(summary["total_size_bytes"]),
            "lines_total": int(summary["lines_total"]),
            "lines_by_kind": dict(summary["lines_by_kind"]),
            "lines_by_language": dict(summary["lines_by_language"]),
        }

    all_records = iter_all_records(
        tracked_records, untracked_records, local_only_records
    )
    overall_summary = summarize_records([r for _, r in all_records])
    content_tokens_total = total_content_tokens(all_records)
    have_tiktoken = _get_token_encoder(token_model) is not None

    payload: Dict[str, object] = {
        "generator": generator_name(),
        "repository": repo.name,
        "root": str(repo),
        "report_created_utc": report_created,
        "git": {
            "has_git": has_git,
            "branch": branch_name if has_git else None,
            "head_commit": head_commit if has_git else None,
            "head_commit_datetime_utc": head_commit_datetime if has_git else None,
            "roots": [
                {
                    "label": render_repo_label(root, repo),
                    "kind": kind,
                    "absolute_path": str(root),
                    "tracked_count": sum(
                        1 for _, r in tracked_records if r.get("git_origin") == root
                    ),
                    "untracked_count": sum(
                        1 for _, r in untracked_records if r.get("git_origin") == root
                    ),
                }
                for root, kind in (git_roots or [])
            ],
            "no_repo_count": len(local_only_records),
        },
        "summary": _serialize_summary(overall_summary),
        "tokens": {
            "content_total_estimate": content_tokens_total,
            "tokenizer_model": token_model,
            "exact": have_tiktoken,
        },
    }
    # Always emit all three buckets so consumers don't have to special-case
    # workspace mode (where files in tracked AND local_only coexist).
    payload["tracked"] = [_serialize_record(p, r) for p, r in tracked_records]
    payload["untracked"] = [_serialize_record(p, r) for p, r in untracked_records]
    payload["local_only"] = [_serialize_record(p, r) for p, r in local_only_records]

    output_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def export_odt(
    repo: Path,
    tracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    untracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    local_only_records: Sequence[Tuple[Path, Dict[str, object]]],
    output_path: Path,
    has_git: bool,
    branch_name: str,
    head_commit: str,
    head_commit_datetime: str,
    report_created: str,
    sections: Optional[Set[str]] = None,
    token_model: str = "gpt-4o",
    git_roots: Optional[Sequence[Tuple[Path, str]]] = None,
) -> None:
    try:
        from odf.opendocument import OpenDocumentText
        from odf.text import H, P
        from odf.style import Style, TextProperties, ParagraphProperties
        from odf.dc import Creator
        from odf.meta import InitialCreator
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency for ODT export. Install with: py -m pip install odfpy"
        ) from exc

    if sections is None:
        sections = set(DEFAULT_SECTIONS)

    doc = OpenDocumentText()
    doc.meta.addElement(InitialCreator(text=generator_name()))
    doc.meta.addElement(Creator(text=generator_name()))

    code_style = Style(name="Code", family="paragraph")
    code_style.addElement(TextProperties(fontfamily="Consolas", fontsize="9pt"))
    code_style.addElement(ParagraphProperties(marginbottom="0.05in", margintop="0in"))
    doc.styles.addElement(code_style)

    meta_style = Style(name="Meta", family="paragraph")
    meta_style.addElement(TextProperties(fontfamily="Consolas", fontsize="8pt"))
    meta_style.addElement(ParagraphProperties(marginbottom="0in", margintop="0in"))
    doc.styles.addElement(meta_style)

    def add_h(level: int, text: str) -> None:
        doc.text.addElement(H(outlinelevel=level, text=text))

    def add_p(text: str = "") -> None:
        doc.text.addElement(P(text=text))

    def add_meta(text: str) -> None:
        doc.text.addElement(P(stylename="Meta", text=text))

    def add_code(text: str) -> None:
        for line in text.splitlines() or [""]:
            doc.text.addElement(P(stylename="Code", text=line))

    content_tokens_total = total_content_tokens(
        iter_all_records(tracked_records, untracked_records, local_only_records)
    )

    add_h(1, f"Repository export: {repo.name}")
    add_p(f"Root folder: {repo}")
    add_p(f"Report created (UTC): {report_created}")
    add_p(
        f"Content tokens (estimate, tokenizer {token_model}): "
        f"~{content_tokens_total:,}"
    )
    if has_git:
        add_p(f"Git repo detected: yes")
        add_p(f"Git branch: {branch_name}")
        add_p(f"HEAD commit: {head_commit}")
        add_p(f"HEAD commit date (UTC): {head_commit_datetime}")
    else:
        add_p("Git repo detected: no")
    add_p(format_git_roots_summary(
        git_roots or [], repo,
        tracked_records, untracked_records, local_only_records,
    ))

    # Unified view; see export_docx / export_md for the workspace-mode
    # rationale (no-repo files coexist with tracked/untracked subrepos).
    all_records = iter_all_records(tracked_records, untracked_records, local_only_records)
    overall_summary = summarize_records([r for _, r in all_records])
    tracked_summary = (
        summarize_records([r for _, r in tracked_records]) if tracked_records else None
    )
    untracked_summary = (
        summarize_records([r for _, r in untracked_records]) if untracked_records else None
    )

    if "glance" in sections:
        add_h(1, "At a glance")
        add_p(
            f"Files: {overall_summary['file_count']} "
            f"({overall_summary['text_count']} text, {overall_summary['binary_count']} binary). "
            f"Total size: {format_file_size_kb(int(overall_summary['total_size_bytes']))}."
        )
        if has_git and tracked_summary and untracked_summary:
            add_p(
                f"Git split: tracked {tracked_summary['file_count']} files / "
                f"{format_file_size_kb(int(tracked_summary['total_size_bytes']))}; "
                f"untracked {untracked_summary['file_count']} files / "
                f"{format_file_size_kb(int(untracked_summary['total_size_bytes']))}."
            )
        add_p(f"Total lines (text files): {int(overall_summary['lines_total']):,}")
        for kind in ("code", "docs", "data", "config", "other"):
            n = int(overall_summary["lines_by_kind"].get(kind, 0))
            if n > 0:
                add_p(f"  {kind}: {n:,} lines")
        top_langs = sorted(
            overall_summary["lines_by_language"].items(),
            key=lambda kv: -int(kv[1]),
        )[:6]
        if top_langs:
            add_p("Top languages by lines:")
            for lang, n in top_langs:
                if int(n) > 0:
                    add_p(f"  {lang}: {int(n):,}")
        if all_records:
            sorted_by_lines = sorted(
                all_records,
                key=lambda item: int(item[1].get("line_count") or 0),
                reverse=True,
            )
            path0, record0 = sorted_by_lines[0]
            n0 = record0.get("line_count")
            if isinstance(n0, int) and n0 > 0:
                add_p(
                    f"Largest file: {path0.relative_to(repo).as_posix()} "
                    f"({n0:,} lines, {record0['size_human']})"
                )
        excerpt = extract_readme_excerpt(all_records, repo)
        if excerpt:
            add_p("README excerpt:")
            for line in excerpt.splitlines():
                add_p(line)

    if "recent" in sections and all_records:
        add_h(1, "Recent files")
        top_records = sorted(all_records, key=lambda item: str(item[1]["mtime_utc"]), reverse=True)[:5]
        for path, record in top_records:
            add_p(f"- {path.relative_to(repo).as_posix()} ({record['mtime_utc']})")

    if "architecture" in sections:
        add_h(1, "Architecture")
        if has_git:
            grouped = group_records_by_directory(repo, list(tracked_records) + list(untracked_records))
        else:
            grouped = group_records_by_directory(repo, local_only_records)
        for bucket, bucket_records in grouped:
            add_p(f"- {bucket} — {describe_directory_bucket(bucket_records)}")

    if "entry_points" in sections:
        add_h(1, "Entry points")
        ranked = collect_ranked_entry_points(all_records, repo)
        if not ranked:
            add_p("(no strong entry-point signals detected)")
        else:
            seen: Set[Tuple[str, str]] = set()
            for _, rel, label, confidence in ranked:
                key = (rel, label)
                if key in seen:
                    continue
                seen.add(key)
                add_p(f"- {rel} [{confidence}] — {label}")

    if "trace" in sections:
        add_h(1, "Entry-point call trace")
        traces = build_entry_point_trace(all_records, repo)
        if not traces:
            add_p("(no traceable entry points)")
        else:
            add_p(
                "Static call sketch — project-defined callees only, library calls omitted. "
                "Two levels deep, regex-based, best-effort."
            )
            for trace in traces:
                add_p(f"- {trace['rel']} :: {trace['seed']}() [{trace['confidence']}] — {trace['label']}")
                for child in trace.get("calls", []) or []:
                    suffix = f" → {child['file']}" if child.get("file") else ""
                    add_p(f"    → {child['name']}(){suffix}")
                    for grand in child.get("calls", []) or []:
                        add_p(f"        → {grand}()")

    if "core" in sections:
        add_h(1, "Core files")
        ranked2: List[Tuple[int, str, str, int, int, int]] = []
        for path, record in all_records:
            rel = path.relative_to(repo).as_posix()
            role = str(record.get("role_hint", ""))
            line_count = int(record.get("line_count") or 0)
            used_by = len(record.get("used_by", []))
            deps = len(record.get("dependencies_resolved", []))
            bonus = 0
            if role == "code":
                bonus += 20
            if role in {"test", "docs", "build", "asset", "config"}:
                bonus -= 50
            score = (line_count // 20) + (used_by * 8) + (deps * 2) + bonus
            ranked2.append((score, rel, role, line_count, used_by, deps))
        for _, rel, role, line_count, used_by, deps in sorted(ranked2, key=lambda x: (-x[0], x[1]))[:8]:
            add_p(f"- {rel} [{role}] — {line_count} lines, used_by: {used_by}, deps: {deps}")

    if "toc" in sections:
        add_h(1, "Table of contents")
        multi_root = bool(git_roots and len(git_roots) > 1)
        for path, record in iter_all_records(
            tracked_records, untracked_records, local_only_records
        ):
            rel = path.relative_to(repo).as_posix()
            line_part = (f", {record['line_count']} lines"
                         if record.get("line_count") is not None else "")
            tok = int(record.get("tokens_estimate", 0))
            tok_part = f", ~{tok:,} tokens" if tok > 0 else ""
            marker = render_toc_marker(record, repo, multi_root)
            add_meta(
                f"[{marker}] {rel} ({record['size_human']}{line_part}"
                f"{tok_part}, {record['mtime_utc']})"
            )

    if "entries" in sections:
        # Heading per non-empty bucket — workspace mode populates all three.
        groups: List[Tuple[str, Sequence[Tuple[Path, Dict[str, object]]]]] = []
        if tracked_records:
            groups.append(("Tracked file entries", tracked_records))
        if untracked_records:
            groups.append(("Untracked file entries", untracked_records))
        if local_only_records:
            title = "No-repo file entries" if has_git else "File entries"
            groups.append((title, local_only_records))
        for section_title, records_seq in groups:
            add_h(1, section_title)
            if not records_seq:
                add_p("(none)")
                continue
            grouped = group_records_by_directory(repo, records_seq)
            for bucket, bucket_records in grouped:
                add_h(2, bucket)
                for path, record in bucket_records:
                    rel = path.relative_to(repo).as_posix()
                    add_h(3, build_file_heading(rel, record))
                    add_meta("--- FILE ---")
                    add_meta(f"path: {rel}")
                    add_meta(f"type: {record['type']}")
                    add_meta(f"role: {record['role_hint']}")
                    add_meta(f"size: {record['size_human']}")
                    add_meta(f"mtime: {record['mtime_utc']}")
                    add_meta(f"id: sha1:{record['quick_hash']}")
                    if record.get("line_count") is not None:
                        add_meta(f"lines: {record['line_count']}")
                    complexity = record.get("complexity", {})
                    classes = int(complexity.get("classes", 0))
                    structs = int(complexity.get("structs", 0))
                    functions_known = bool(complexity.get("functions_known", False))
                    functions = int(complexity.get("functions", 0)) if functions_known else 0
                    if classes or structs or functions:
                        bits: List[str] = []
                        if functions:
                            bits.append(f"functions={functions}")
                        if classes:
                            bits.append(f"classes={classes}")
                        if structs:
                            bits.append(f"structs={structs}")
                        add_meta("decls: " + ", ".join(bits))
                    deps_list = record.get("dependencies_resolved", [])
                    add_meta(f"deps: {', '.join(deps_list) if deps_list else '(none)'}")
                    used_by_list = record.get("used_by", [])
                    add_meta(f"used_by: {', '.join(used_by_list) if used_by_list else '(none)'}")
                    signals_list = [f"{label} [{confidence}]" for label, confidence in record.get("entry_signals", [])]
                    add_meta(f"entry_signals: {', '.join(signals_list) if signals_list else '(none)'}")
                    if record["content_included"]:
                        add_meta("--- BEGIN CONTENT ---")
                        add_code(sanitize_xml_compatible_text(str(record["content_or_note"])))
                        add_meta("--- END FILE ---")
                    else:
                        add_meta(f"omission: {record['omission_reason'] or 'binary or non-text file'}")
                        add_meta(f"note: {record['content_or_note']}")
                        add_meta("--- END FILE ---")

    doc.save(output_path)


def export_ods(
    repo: Path,
    tracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    untracked_records: Sequence[Tuple[Path, Dict[str, object]]],
    local_only_records: Sequence[Tuple[Path, Dict[str, object]]],
    output_path: Path,
    has_git: bool,
    branch_name: str,
    head_commit: str,
    head_commit_datetime: str,
    report_created: str,
    token_model: str = "gpt-4o",
    git_roots: Optional[Sequence[Tuple[Path, str]]] = None,
) -> None:
    try:
        from odf.opendocument import OpenDocumentSpreadsheet
        from odf.table import Table, TableRow, TableCell, TableColumn
        from odf.text import P
        from odf.dc import Creator
        from odf.meta import InitialCreator
        from odf.style import (
            Style,
            TableColumnProperties,
            TableRowProperties,
            TableCellProperties,
            ParagraphProperties,
            TextProperties,
        )
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency for ODS export. Install with: py -m pip install odfpy"
        ) from exc

    headers = [
        "section", "git_origin",
        "directory_group", "file_path", "type", "kind", "role",
        "size_bytes", "size_human", "mtime_utc", "quick_hash", "is_text",
        "line_count", "tokens_estimate",
        "functions_est", "classes_est", "structs_est",
        "dependencies", "used_by", "entry_signals",
        "content_included", "omission_reason", "content_or_note",
    ]
    # Column widths in mm, one per header. Mirrors the xlsx layout: file_path
    # and the free-text columns get extra room; numeric/flag columns stay tight.
    column_widths_mm = [
        25, 28, 35, 100, 28, 22, 22,
        22, 28, 38, 24, 16,
        20, 22,
        20, 20, 20,
        60, 60, 50,
        24, 50, 200,
    ]

    doc = OpenDocumentSpreadsheet()
    doc.meta.addElement(InitialCreator(text=generator_name()))
    doc.meta.addElement(Creator(text=generator_name()))

    # Top-aligned cells with wrap so multi-line content flows downward inside
    # a fixed-height row rather than horizontally.
    cell_style = Style(name="WrapTop", family="table-cell")
    cell_style.addElement(TableCellProperties(verticalalign="top", wrapoption="wrap"))
    doc.automaticstyles.addElement(cell_style)

    header_cell_style = Style(name="HeaderCell", family="table-cell")
    header_cell_style.addElement(TableCellProperties(verticalalign="top", wrapoption="wrap"))
    doc.automaticstyles.addElement(header_cell_style)
    header_para_style = Style(name="HeaderPara", family="paragraph")
    header_para_style.addElement(TextProperties(fontweight="bold"))
    header_para_style.addElement(ParagraphProperties())
    doc.automaticstyles.addElement(header_para_style)

    # Fixed row heights so 'content_or_note' (which can hold a full source
    # file) does not stretch a single row to the full screen height.
    # use-optimal-row-height="false" stops LibreOffice from auto-fitting.
    header_row_style = Style(name="HeaderRow", family="table-row")
    header_row_style.addElement(
        TableRowProperties(rowheight="6mm", useoptimalrowheight="false")
    )
    doc.automaticstyles.addElement(header_row_style)
    data_row_style = Style(name="DataRow", family="table-row")
    data_row_style.addElement(
        TableRowProperties(rowheight="8mm", useoptimalrowheight="false")
    )
    doc.automaticstyles.addElement(data_row_style)

    # One column-style per width so each column gets its own size.
    column_style_names: List[str] = []
    for idx, mm in enumerate(column_widths_mm):
        name = f"Col{idx}"
        col_style = Style(name=name, family="table-column")
        col_style.addElement(TableColumnProperties(columnwidth=f"{mm}mm"))
        doc.automaticstyles.addElement(col_style)
        column_style_names.append(name)

    table = Table(name="Repository export")
    for name in column_style_names:
        table.addElement(TableColumn(stylename=name))

    def make_row(
        values: Sequence[object],
        *,
        row_style: Optional[str] = None,
        cell_style_name: Optional[str] = None,
        para_style_name: Optional[str] = None,
    ) -> TableRow:
        row = TableRow(stylename=row_style) if row_style else TableRow()
        for v in values:
            cell_kwargs: Dict[str, str] = {}
            if cell_style_name:
                cell_kwargs["stylename"] = cell_style_name
            if isinstance(v, bool):
                cell = TableCell(
                    valuetype="boolean",
                    booleanvalue="true" if v else "false",
                    **cell_kwargs,
                )
                text = "TRUE" if v else "FALSE"
            elif isinstance(v, (int, float)):
                cell = TableCell(valuetype="float", value=str(v), **cell_kwargs)
                text = str(v)
            else:
                cell = TableCell(valuetype="string", **cell_kwargs)
                text = str(v) if v is not None else ""
            if para_style_name:
                cell.addElement(P(stylename=para_style_name, text=text))
            else:
                cell.addElement(P(text=text))
            row.addElement(cell)
        return row

    table.addElement(make_row(
        headers,
        row_style="HeaderRow",
        cell_style_name="HeaderCell",
        para_style_name="HeaderPara",
    ))

    # Iterate all three buckets unconditionally — see export_xlsx.
    groups = [
        ("tracked", tracked_records),
        ("untracked", untracked_records),
        ("local", local_only_records),
    ]

    for section_name, records in groups:
        for path, record in records:
            rel = path.relative_to(repo).as_posix()
            directory_group = infer_directory_bucket(rel)
            complexity = record.get("complexity", {})
            entry_signals_text = ", ".join(
                f"{label} [{confidence}]" for label, confidence in record.get("entry_signals", [])
            )
            functions_cell = (
                complexity.get("functions", 0) if complexity.get("functions_known") else ""
            )
            line_cell = record["line_count"] if record["line_count"] is not None else ""
            origin = record.get("git_origin")
            git_origin_label = render_repo_label(origin, repo) if origin else ""
            table.addElement(make_row(
                [
                    section_name,
                    git_origin_label,
                    directory_group,
                    rel,
                    record["type"],
                    record.get("kind", ""),
                    record["role_hint"],
                    record["size_bytes"],
                    record["size_human"],
                    record["mtime_utc"],
                    record["quick_hash"],
                    record["is_text"],
                    line_cell,
                    int(record.get("tokens_estimate", 0)),
                    functions_cell,
                    complexity.get("classes", 0),
                    complexity.get("structs", 0),
                    ", ".join(record.get("dependencies_resolved", [])),
                    ", ".join(record.get("used_by", [])),
                    entry_signals_text,
                    record["content_included"],
                    record["omission_reason"],
                    sanitize_xml_compatible_text(str(record["content_or_note"])),
                ],
                row_style="DataRow",
                cell_style_name="WrapTop",
            ))

    doc.spreadsheet.addElement(table)

    # Summary sheet: two narrow label/value columns.
    summary_col_a = Style(name="SumColA", family="table-column")
    summary_col_a.addElement(TableColumnProperties(columnwidth="55mm"))
    doc.automaticstyles.addElement(summary_col_a)
    summary_col_b = Style(name="SumColB", family="table-column")
    summary_col_b.addElement(TableColumnProperties(columnwidth="120mm"))
    doc.automaticstyles.addElement(summary_col_b)

    content_tokens_total = total_content_tokens(
        iter_all_records(tracked_records, untracked_records, local_only_records)
    )

    summary_rows: List[Tuple[object, object]] = [
        ("Repository", repo.name),
        ("Root folder", str(repo)),
        ("Report created (UTC)", report_created),
        ("Git repo detected", "yes" if has_git else "no"),
        ("Branch", branch_name if has_git else "n/a"),
        ("HEAD commit", head_commit if has_git else "n/a"),
        ("HEAD commit date (UTC)", head_commit_datetime if has_git else "n/a"),
        ("Content tokens (estimate)", content_tokens_total),
        ("Tokenizer model", token_model),
        ("Git roots summary", format_git_roots_summary(
            git_roots or [], repo,
            tracked_records, untracked_records, local_only_records,
        )),
    ]
    for root, kind in (git_roots or []):
        t = sum(1 for _, r in tracked_records if r.get("git_origin") == root)
        u = sum(1 for _, r in untracked_records if r.get("git_origin") == root)
        summary_rows.append((
            f"  · {render_repo_label(root, repo)}",
            f"{kind}; tracked={t}, untracked={u}",
        ))
    if local_only_records:
        summary_rows.append(("  · no-repo files", len(local_only_records)))

    summary_table = Table(name="Summary")
    summary_table.addElement(TableColumn(stylename="SumColA"))
    summary_table.addElement(TableColumn(stylename="SumColB"))
    for label, value in summary_rows:
        summary_table.addElement(make_row([label, value], row_style="DataRow"))
    doc.spreadsheet.addElement(summary_table)

    doc.save(output_path)


def apply_strip_comments(
    records: Sequence[Tuple[Path, Dict[str, object]]],
) -> int:
    """Apply strip_comments_only to every record's 'content_or_note'.
    Returns the number of records whose content was modified."""
    n = 0
    for path, record in records:
        if not record.get("content_included"):
            continue
        content = record.get("content_or_note")
        if not isinstance(content, str) or not content:
            continue
        suffix = path.suffix.lower()
        stripped = strip_comments_only(content, suffix)
        if stripped != content:
            record["content_or_note"] = stripped
            n += 1
    return n


def _token_count_for(text: str, model: str) -> int:
    """Shared token estimator. Uses tiktoken when available, falls back to a
    4-char/token rough estimate otherwise."""
    t = count_tokens(text, model)
    return t if t is not None else estimate_tokens_rough(text)


def annotate_record_tokens(
    records: Sequence[Tuple[Path, Dict[str, object]]],
    model: str,
) -> int:
    """Add a 'tokens_estimate' field to every record reflecting the tokens
    consumed by its content_or_note. Returns the total across all records.

    Called once before md/json export so the per-file metadata, the toc,
    and the budget logic all share the same numbers (no double tokenize)."""
    total = 0
    for _, record in records:
        content = str(record.get("content_or_note", ""))
        n = _token_count_for(content, model)
        record["tokens_estimate"] = n
        total += n
    return total


def apply_token_budget(
    records: Sequence[Tuple[Path, Dict[str, object]]],
    budget: int,
    model: str,
) -> Tuple[int, int, int, int]:
    """If the combined content tokens exceed `budget`, replace the
    content_or_note of the lowest-ranked records with a trim marker until
    the budget is satisfied. Ranking is by used_by then size (see
    rank_records_for_pruning), and trimming starts from the bottom.

    Reads pre-computed 'tokens_estimate' from each record (call
    annotate_record_tokens first); updates the field when content is
    swapped for the trim marker.

    Returns (kept_records, trimmed_records, tokens_before, tokens_after)."""
    ranked = rank_records_for_pruning(records)  # high-value first
    total = sum(int(record.get("tokens_estimate", 0)) for _, record in ranked)
    before = total

    # Trim from the bottom of the ranking (lowest value) until under budget.
    trimmed = 0
    for path, record in reversed(ranked):
        if total <= budget:
            break
        if not record.get("content_included"):
            continue
        was = int(record.get("tokens_estimate", 0))
        if was <= 0:
            continue
        record["content_included"] = False
        record["omission_reason"] = (
            record.get("omission_reason") or "trimmed to fit --token-budget"
        )
        record["content_or_note"] = (
            f"[trimmed to fit --token-budget {budget}; "
            "see source file for full content]"
        )
        new = _token_count_for(str(record["content_or_note"]), model)
        record["tokens_estimate"] = new
        total = total - was + new
        trimmed += 1
    return (len(ranked) - trimmed, trimmed, before, total)


def main() -> int:
    args = parse_args()

    # --- repo resolution: either local path or clone --remote into tempdir
    remote_tempdir: Optional["tempfile.TemporaryDirectory"] = None
    if args.remote:
        if args.repo:
            print(
                "ERROR: pass either a local path or --remote, not both.",
                file=sys.stderr,
            )
            return 1
        try:
            repo, remote_tempdir = clone_remote_to_tempdir(args.remote)
            print(f"Cloned {args.remote} -> {repo}", file=sys.stderr)
        except RuntimeError as exc:
            print(f"ERROR: {exc}", file=sys.stderr)
            return 1
    else:
        if not args.repo:
            print(
                "ERROR: provide a repo path or use --remote owner/repo.",
                file=sys.stderr,
            )
            return 1
        repo = Path(args.repo).expanduser().resolve()
        if not repo.exists() or not repo.is_dir():
            print(f"ERROR: Folder does not exist: {repo}", file=sys.stderr)
            return 1

    try:
        return _run_export(args, repo)
    finally:
        if remote_tempdir is not None:
            remote_tempdir.cleanup()


def _run_export(args: argparse.Namespace, repo: Path) -> int:
    try:
        sections = resolve_sections(args.sections)
    except ValueError as exc:
        print(f"ERROR: --sections: {exc}", file=sys.stderr)
        return 1

    allowed_extensions = None
    if args.ext is not None and len(args.ext) > 0:
        allowed_extensions = {
            ext.lower() if ext.startswith(".") else f".{ext.lower()}"
            for ext in args.ext
        }

    exclude_dirs = set(DEFAULT_EXCLUDE_DIRS) | set(args.exclude_dir)
    exclude_patterns = set(DEFAULT_EXCLUDE_PATTERNS) | set(args.exclude_pattern)

    files = collect_all_files(
        repo=repo,
        allowed_extensions=allowed_extensions,
        exclude_dirs=exclude_dirs,
        exclude_patterns=exclude_patterns,
    )

    # Discover any git work tree(s) that touch repo: repo itself, an ancestor,
    # or immediate-child subfolders. Empty list = pure local-only workspace.
    roots = discover_git_roots(repo)
    if roots:
        kinds = ", ".join(
            f"{render_repo_label(root, repo)} ({kind})" for root, kind in roots
        )
        print(f"Git repos detected: {kinds}.", file=sys.stderr)

    # .gitignore filter (default on, suppressed by --no-gitignore). Loops
    # over every discovered root; files outside any root are passed through.
    if not args.no_gitignore and roots:
        before_count = len(files)
        files = filter_by_gitignore_for_roots(roots, files)
        dropped = before_count - len(files)
        if dropped > 0:
            print(f".gitignore filtered out {dropped} file(s).", file=sys.stderr)

    # --since <ref> restricts to files changed since that revision in any
    # discovered repo. Files outside any repo are excluded.
    if args.since:
        if not roots:
            print(
                "ERROR: --since requires a git repository (none found at, "
                f"above, or below {repo}).",
                file=sys.stderr,
            )
            return 1
        try:
            changed_abs = get_changed_paths_for_roots(roots, args.since)
        except RuntimeError as exc:
            print(f"ERROR: {exc}", file=sys.stderr)
            return 1
        files = [p for p in files if p in changed_abs]
        print(
            f"--since {args.since}: kept {len(files)} changed file(s).",
            file=sys.stderr,
        )

    if not files:
        print("No files matched the selection criteria.", file=sys.stderr)
        return 2

    has_git = bool(roots)
    # Pick a representative root for the legacy header fields. When there are
    # multiple, prefer the first child (workspace mode); the multi-repo
    # summary line above already lists everything.
    primary_root = roots[0][0] if roots else None
    branch_name = get_git_branch_name(primary_root) if primary_root else "n/a"
    head_commit = get_git_head_commit(primary_root) if primary_root else "n/a"
    head_commit_datetime = (
        get_git_head_commit_datetime_utc(primary_root) if primary_root else "n/a"
    )
    report_created = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    # Per-root tracked sets, fetched once. Used to tag each record's
    # git_tracked field during build_records_for_files.
    tracked_per_root: Dict[Path, Set[str]] = {}
    for root, _ in roots:
        try:
            relpaths = get_git_tracked_relpaths(root)
            tracked_per_root[root] = {p.as_posix() for p in relpaths}
        except RuntimeError as exc:
            print(f"WARNING: cannot list tracked files in {root}: {exc}",
                  file=sys.stderr)
            tracked_per_root[root] = set()

    # Single pass: build every record at once, tagging git_origin and
    # git_tracked inline. The three legacy buckets are derived afterward.
    all_records = build_records_for_files(
        repo=repo,
        files=files,
        preferred_encoding=args.encoding,
        max_text_file_kb=args.max_text_file_kb,
        include_no_extension=args.include_no_extension,
        roots=roots,
        tracked_per_root=tracked_per_root,
    )

    # Derive the legacy three-bucket view for the exporter APIs that still
    # take them. tracked = in some repo AND tracked. untracked = in some
    # repo but not tracked. local_only = outside any repo (the new
    # "no-repo" bucket; was "non-git" in the old code).
    tracked_records: List[Tuple[Path, Dict[str, object]]] = []
    untracked_records: List[Tuple[Path, Dict[str, object]]] = []
    local_only_records: List[Tuple[Path, Dict[str, object]]] = []
    for path, record in all_records:
        if record.get("git_origin") is None:
            local_only_records.append((path, record))
        elif record.get("git_tracked"):
            tracked_records.append((path, record))
        else:
            untracked_records.append((path, record))

    enrich_cross_file_metadata(repo, all_records)

    # --strip-comments shrinks code content before any rendering or token
    # accounting so the token-budget calculation reflects the trimmed text.
    if args.strip_comments:
        modified = apply_strip_comments(all_records)
        print(f"--strip-comments: trimmed {modified} file(s).", file=sys.stderr)

    fmt = resolve_format(args.format, args.output)

    # Annotate every record with a per-file token estimate. Cheap with
    # tiktoken (~1 ms per kB of text); useful in every format because a
    # token count is the only size unit that's tokenizer-comparable across
    # languages (10 kB of Python ~= 4 k tokens; 10 kB of Java ~= 2.7 k).
    # md/json budget logic reuses the same numbers (no double tokenize).
    annotate_record_tokens(all_records, args.token_model)

    # --token-budget is only meaningful for text outputs (md/json). For
    # binary formats it has no effect; warn rather than silently ignore.
    if args.token_budget is not None:
        if fmt in ("md", "json"):
            kept, trimmed, before, after = apply_token_budget(
                all_records, args.token_budget, args.token_model
            )
            if trimmed > 0:
                print(
                    f"--token-budget {args.token_budget}: trimmed {trimmed} "
                    f"file(s); ~{before:,} -> ~{after:,} content tokens.",
                    file=sys.stderr,
                )
            else:
                print(
                    f"--token-budget {args.token_budget}: under budget "
                    f"(~{before:,} content tokens, no trimming).",
                    file=sys.stderr,
                )
        else:
            print(
                f"--token-budget ignored: only applies to md/json (got {fmt}).",
                file=sys.stderr,
            )

    if args.output:
        output_path = Path(args.output).expanduser().resolve()
    else:
        output_path = repo.parent / f"{repo.name}_code_export.{fmt}"

    output_path.parent.mkdir(parents=True, exist_ok=True)

    common_kwargs = dict(
        repo=repo,
        tracked_records=tracked_records,
        untracked_records=untracked_records,
        local_only_records=local_only_records,
        output_path=output_path,
        has_git=has_git,
        branch_name=branch_name,
        head_commit=head_commit,
        head_commit_datetime=head_commit_datetime,
        report_created=report_created,
        git_roots=roots,
    )

    # git_roots is already part of common_kwargs (every exporter accepts it).
    try:
        if fmt == "docx":
            export_docx(sections=sections, token_model=args.token_model,
                        **common_kwargs)
        elif fmt == "xlsx":
            export_xlsx(token_model=args.token_model, **common_kwargs)
        elif fmt == "md":
            export_md(sections=sections, token_model=args.token_model,
                      **common_kwargs)
        elif fmt == "json":
            export_json(token_model=args.token_model, **common_kwargs)
        elif fmt == "odt":
            export_odt(sections=sections, token_model=args.token_model,
                       **common_kwargs)
        elif fmt == "ods":
            export_ods(token_model=args.token_model, **common_kwargs)
        else:
            print(f"ERROR: unknown format: {fmt}", file=sys.stderr)
            return 1
    except RuntimeError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 3

    # --clipboard for text outputs only.
    if args.clipboard:
        if fmt in ("md", "json"):
            try:
                payload = output_path.read_text(encoding="utf-8")
            except OSError as exc:
                print(f"WARNING: could not re-read output for clipboard: {exc}",
                      file=sys.stderr)
            else:
                if copy_text_to_clipboard(payload):
                    print(f"Copied {len(payload):,} chars to clipboard.",
                          file=sys.stderr)
                else:
                    print(
                        "WARNING: clipboard copy failed (install pyperclip or "
                        "ensure clip.exe / pbcopy / wl-copy / xclip is available).",
                        file=sys.stderr,
                    )
        else:
            print(
                f"--clipboard ignored: only applies to md/json (got {fmt}).",
                file=sys.stderr,
            )

    # Unified summary — works for single-repo, workspace (multi subrepo),
    # ancestor, and no-repo cases without three separate code paths.
    tracked_count = len(tracked_records)
    untracked_count = len(untracked_records)
    no_repo_count = len(local_only_records)
    total_count = tracked_count + untracked_count + no_repo_count
    total_size = sum(int(record["size_bytes"]) for _, record in all_records)
    print(f"Export created: {output_path}")
    if has_git:
        print(f"Branch: {branch_name}  HEAD: {head_commit}")
    if tracked_count or untracked_count:
        print(f"Tracked: {tracked_count}  Untracked: {untracked_count}", end="")
        if no_repo_count:
            print(f"  No-repo: {no_repo_count}", end="")
        print()
    else:
        print(f"Files: {total_count} (no git)")
    print(f"Total files: {total_count}, total size: {format_file_size_kb(total_size)}")

    # Final token report on md/json. Helps users decide whether to add or
    # raise --token-budget. Cheap if tiktoken isn't installed (rough estimate).
    if fmt in ("md", "json"):
        try:
            payload = output_path.read_text(encoding="utf-8")
        except OSError:
            payload = ""
        if payload:
            t = count_tokens(payload, args.token_model)
            label = "tokens" if t is not None else "tokens (rough estimate)"
            if t is None:
                t = estimate_tokens_rough(payload)
            print(f"Output {label}: ~{t:,}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
